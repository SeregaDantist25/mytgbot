# -*- coding: utf-8 -*-
"""Загрузка Word-шаблонов и безопасная замена плейсхолдеров."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document as DocxDocument


class TemplateService:
    def __init__(self, templates_dir: str | Path):
        self.templates_dir = Path(templates_dir)

    def load(self, filename: str):
        template_path = self.templates_dir / filename
        if not template_path.is_file():
            raise FileNotFoundError(
                f"Шаблон {filename} не найден в {self.templates_dir}"
            )
        return DocxDocument(template_path)

    @staticmethod
    def _merge_runs_with_tag(paragraph, tag: str) -> None:
        """Собрать разбитый Word-ом тег в первый run параграфа."""
        full_text = paragraph.text
        if tag not in full_text:
            return
        if any(tag in run.text for run in paragraph.runs):
            return
        if not paragraph.runs:
            return
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""

    def replace_placeholders(self, document, placeholders: dict):
        """Заменить теги ``{{key}}`` в тексте и таблицах документа."""
        def replace_in_paragraph(paragraph):
            for key, value in placeholders.items():
                tag = f"{{{{{key}}}}}"
                if tag not in paragraph.text:
                    continue
                self._merge_runs_with_tag(paragraph, tag)
                for run in paragraph.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, str(value))

        for paragraph in document.paragraphs:
            replace_in_paragraph(paragraph)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)
        return document


TEMPLATES_DIR = Path(os.getenv("TEMPLATES_DIR", "templates"))
service = TemplateService(TEMPLATES_DIR)


def load_template(filename: str):
    return service.load(filename)


def replace_placeholders(document, placeholders: dict):
    return service.replace_placeholders(document, placeholders)
