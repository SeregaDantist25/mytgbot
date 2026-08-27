# -*- coding: utf-8 -*-
"""
Прочая вспомогательная бизнес-логика, вынесенная из bot.py.

Содержит:
- загрузку данных из JSON (чек-листы, суда, компании)
- класс PumpDatabase (база знаний по насосам)
- работу с шаблонами Word и плейсхолдерами
- счётчики документов (counters.db)
- персистентное состояние диалога (chat_state.json)
- git-хелперы (автокоммит/автопуш конфигов)
- добавление судов и компаний
- детекцию и парсинг текста (оборудование, судно, дефекты, зазоры, АВР)
- генерацию объёма работ и таблиц дефектации
- функции ORM (загрузка ремонтной ведомости, роли, версионирование)
"""

import os
import re
import json
import time
import sqlite3
import subprocess
import logging
from datetime import datetime
from typing import Optional

from docx import Document as DocxDocument

from models import (
    SessionLocal,
    User,
    Ship,
    RepairStatement,
    StatementItem,
)
from services.document_service import (
    approve_document,
    archive_document,
    delete_document,
)
from services.chat_state_service import get_chat_state, set_chat_state

logger = logging.getLogger(__name__)

# --- Пути к файлам ---
TEMPLATES_DIR = os.getenv("TEMPLATES_DIR", "templates")
DATA_DIR = os.getenv("DATA_DIR", "data")
CHECKLISTS_FILE = os.path.join(DATA_DIR, "checklists.json")
COUNTERS_FILE = os.path.join(DATA_DIR, "counters.json")
SHIPS_FILE = os.path.join(DATA_DIR, "ships.json")
COMPANIES_FILE = os.path.join(DATA_DIR, "companies.json")
EMPLOYEES_FILE = os.path.join(DATA_DIR, "employees.json")
COUNTERS_DB = os.path.join(DATA_DIR, "counters.db")

# Полный путь к git (в PATH его может не быть)
_GIT_EXE = r"C:\Program Files\Git\bin\git.exe"


# ============================================================
#  ЗАГРУЗКА ДАННЫХ ИЗ JSON
# ============================================================

def load_checklists() -> dict:
    """Загружает чек-листы из data/checklists.json."""
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    if not os.path.exists(CHECKLISTS_FILE):
        return {}
    try:
        with open(CHECKLISTS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}


def load_ships() -> dict:
    """Загружает словарь судов из data/ships.json."""
    if not os.path.exists(SHIPS_FILE):
        return {}
    with open(SHIPS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)


def load_employees() -> list:
    """Загружает список сотрудников из data/employees.json.

    Returns:
        Список словарей вида {"name": str, "role": str}.
    """
    if not os.path.exists(EMPLOYEES_FILE):
        return []
    try:
        with open(EMPLOYEES_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data.get("employees", [])
    except Exception:
        return []


def find_employee_role(name: str) -> Optional[str]:
    """Ищет роль сотрудника по ФИО (без учёта регистра).

    Args:
        name: ФИО пользователя.

    Returns:
        Роль сотрудника или None, если ФИО не найдено.
    """
    if not name:
        return None
    normalized = " ".join(name.strip().lower().split())
    for emp in load_employees():
        emp_name = " ".join(str(emp.get("name", "")).strip().lower().split())
        if emp_name == normalized:
            return emp.get("role")
    return None


def load_companies() -> dict:
    """Загружает дефолтные executor/customer/location из data/companies.json."""
    defaults = {
        "executor": "ООО «Новое время»",
        "customer": "АО «Бункерная компания»",
        "location": "Рейд 4ый район, г. Находка",
    }
    if not os.path.exists(COMPANIES_FILE):
        return defaults
    with open(COMPANIES_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    defaults.update(data)
    return defaults


# ============================================================
#  БАЗА ДАННЫХ НАСОСОВ
# ============================================================

class PumpDatabase:
    """База знаний по насосам (чек-листы, зазоры, дефекты)."""

    def __init__(self) -> None:
        self.data = load_checklists()

    def get_pump_types(self) -> list:
        return list(self.data.keys())

    def get_pump_name(self, pump_type: str) -> str:
        return self.data.get(pump_type, {}).get("name", pump_type)

    def get_checklist(self, pump_type: str) -> list:
        return self.data.get(pump_type, {}).get("items", [])

    def get_clearances(self, pump_type: str, clearance_type: str):
        clearances = self.data.get(pump_type, {}).get("clearances", {})
        return clearances.get(clearance_type)

    def check_clearance(self, pump_type: str, clearance_type: str, measured_value: float) -> dict:
        clearance_data = self.get_clearances(pump_type, clearance_type)
        if not clearance_data:
            return {
                "status": "unknown",
                "message": f"Данные по зазору '{clearance_type}' для '{pump_type}' отсутствуют",
                "action": "Проверьте правильность ввода",
            }

        standard_min = clearance_data.get("min", 0)
        standard_max = clearance_data.get("max", 0)
        unit = clearance_data.get("unit", "мм")

        if "мм/мм" in unit:
            return {
                "status": "info",
                "message": f"📌 Зазор зависит от диаметра: {standard_min}-{standard_max} {unit}",
                "action": "Уточните диаметр для точного расчёта",
            }

        if measured_value < standard_min:
            return {
                "status": "warning",
                "message": f"⚠️ Зазор МЕНЬШЕ нормы: {measured_value} мм (норма: {standard_min}-{standard_max} мм)",
                "action": "Проверьте точность измерения",
            }
        elif measured_value <= standard_max:
            return {
                "status": "ok",
                "message": f"✅ Зазор В НОРМЕ: {measured_value} мм (норма: {standard_min}-{standard_max} мм)",
                "action": "Деталь работоспособна",
            }
        else:
            return {
                "status": "critical",
                "message": f"🔴 Зазор ПРЕВЫШЕН: {measured_value} мм (норма: {standard_min}-{standard_max} мм)",
                "action": "Требуется ремонт",
            }

    def get_common_defects(self, pump_type: str) -> list:
        return self.data.get(pump_type, {}).get("defects", [])

    def get_repair_method(self, pump_type: str, defect_text: str):
        defect_lower = defect_text.lower()
        methods = self.data.get(pump_type, {}).get("repair_methods", {})
        for key, method in methods.items():
            if key in defect_lower:
                return method
        return None


# ============================================================
#  РАБОТА С ШАБЛОНАМИ
# ============================================================

def load_template(filename: str) -> DocxDocument:
    """Загружает шаблон Word из TEMPLATES_DIR."""
    template_path = os.path.join(TEMPLATES_DIR, filename)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон {filename} не найден в {TEMPLATES_DIR}")
    return DocxDocument(template_path)


def _merge_runs_with_tag(paragraph, tag: str) -> None:
    """Склеивает runs параграфа в один, если тег разбит на несколько runs."""
    full_text = paragraph.text
    if tag not in full_text:
        return
    for run in paragraph.runs:
        if tag in run.text:
            return  # тег уже в одном run — нормализация не нужна
    if not paragraph.runs:
        return
    first_run = paragraph.runs[0]
    first_run.text = full_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_placeholders(doc, placeholders: dict) -> DocxDocument:
    """Заменяет плейсхолдеры {{key}} в параграфах и таблицах документа."""
    def _replace_in_paragraph(paragraph):
        for key in placeholders:
            tag = f"{{{{{key}}}}}"
            if tag in paragraph.text:
                _merge_runs_with_tag(paragraph, tag)
                for run in paragraph.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, str(placeholders[key]))

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph)

    return doc


# ============================================================
#  СЧЁТЧИКИ ДОКУМЕНТОВ
# ============================================================

def _init_counters_db() -> None:
    """Создаёт таблицу counters и мигрирует значения из counters.json один раз."""
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    conn = sqlite3.connect(COUNTERS_DB)
    cursor = conn.cursor()
    cursor.execute(
        "CREATE TABLE IF NOT EXISTS counters (doc_type TEXT PRIMARY KEY, value INTEGER)"
    )
    if os.path.exists(COUNTERS_FILE):
        try:
            with open(COUNTERS_FILE, 'r', encoding='utf-8') as f:
                old = json.load(f)
            for doc_type, value in old.items():
                cursor.execute(
                    "INSERT OR IGNORE INTO counters (doc_type, value) VALUES (?, ?)",
                    (doc_type, value),
                )
            conn.commit()
            os.rename(COUNTERS_FILE, COUNTERS_FILE + ".migrated")
        except Exception as e:
            logger.warning(f"Ошибка миграции счётчиков: {e}")
    conn.commit()
    conn.close()


_init_counters_db()


def get_next_number(doc_type: str) -> int:
    """Атомарно инкрементирует счётчик и возвращает новое значение."""
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    conn = sqlite3.connect(COUNTERS_DB)
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE counters SET value = value + 1 WHERE doc_type = ? RETURNING value",
        (doc_type,),
    )
    row = cursor.fetchone()
    if row is None:
        cursor.execute("INSERT INTO counters (doc_type, value) VALUES (?, 1)", (doc_type,))
        conn.commit()
        conn.close()
        return 1
    conn.commit()
    conn.close()
    return row[0]


def get_counter(doc_type: str) -> int:
    """Обратная совместимость: возвращает следующий номер без инкремента."""
    return get_next_number(doc_type)


def update_counter(doc_type: str, new_number: int) -> None:
    """Обратная совместимость: устанавливает счётчик в заданное значение."""
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    conn = sqlite3.connect(COUNTERS_DB)
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO counters (doc_type, value) VALUES (?, ?) "
        "ON CONFLICT(doc_type) DO UPDATE SET value = excluded.value",
        (doc_type, new_number),
    )
    conn.commit()
    conn.close()


# ============================================================
#  GIT: АВТОКОММИТ И АВТОПУШ КОНФИГОВ
# ============================================================

def _git(*args) -> tuple:
    """Выполняет git-команду и возвращает (returncode, output)."""
    try:
        result = subprocess.run(
            [_GIT_EXE] + list(args),
            capture_output=True, text=True, encoding="utf-8", errors="replace",
        )
        return result.returncode, (result.stdout + result.stderr).strip()
    except Exception as e:
        return 1, str(e)


def git_commit_and_push(files: list, message: str) -> tuple:
    """Добавляет файлы, коммитит и пушит. Возвращает (ok, text)."""
    code, out = _git("add", "--", *files)
    if code != 0:
        return False, f"git add: {out}"
    code, out = _git("commit", "-m", message)
    if code != 0 and "nothing to commit" not in out.lower():
        return False, f"git commit: {out}"
    code, out = _git("push", "origin", "master")
    if code != 0:
        return False, f"git push: {out}"
    return True, out


# ============================================================
#  ДОБАВЛЕНИЕ СУДОВ И КОМПАНИЙ (С АВТОПУШЕМ)
# ============================================================

def add_ship(name: str) -> tuple:
    """Добавляет судно в ships.json и пушит. Возвращает (ok, text)."""
    name = name.strip()
    if not name:
        return False, "Пустое название судна."
    ships = load_ships()
    key = name.lower()
    if key in ships:
        return False, f"Судно «{name}» уже есть в списке."
    ships[key] = name
    with open(SHIPS_FILE, 'w', encoding='utf-8') as f:
        json.dump(ships, f, ensure_ascii=False, indent=2)
    ok, text = git_commit_and_push([SHIPS_FILE], f"feat: добавить судно {name}")
    if ok:
        return True, f"✅ Судно «{name}» добавлено и запушено в репозиторий."
    return True, f"✅ Судно «{name}» добавлено в файл, но пуш не удался: {text}"


def add_company(field: str, value: str) -> tuple:
    """Обновляет поле компании (executor/customer/location) и пушит."""
    value = value.strip()
    if not value:
        return False, "Пустое значение."
    companies = load_companies()
    companies[field] = value
    with open(COMPANIES_FILE, 'w', encoding='utf-8') as f:
        json.dump(companies, f, ensure_ascii=False, indent=2)
    ok, text = git_commit_and_push([COMPANIES_FILE], f"feat: обновить {field} в companies.json")
    if ok:
        return True, f"✅ Поле «{field}» обновлено и запушено в репозиторий."
    return True, f"✅ Поле «{field}» обновлено в файле, но пуш не удался: {text}"


# ============================================================
#  ОПРЕДЕЛЕНИЕ ТИПА ОБОРУДОВАНИЯ (ЛОКАЛЬНОЕ)
# ============================================================

def detect_equipment_type(text: str) -> str | None:
    """Определяет тип оборудования локально (без AI)."""
    text_lower = text.lower()
    if any(word in text_lower for word in ["двигател", "дизель", "мотор", "гд"]):
        return "engine"
    elif any(word in text_lower for word in ["компрессор", "компрес"]):
        return "compressor"
    elif any(word in text_lower for word in ["насос", "помп"]):
        return "pump"
    else:
        return None


def ask_for_clarification(equipment: str) -> str:
    """Формирует вопрос для уточнения типа оборудования."""
    return (
        f"🔍 Я нашёл упоминание '{equipment}'. Уточните, это:\n"
        "1️⃣ Насос\n2️⃣ Двигатель\n3️⃣ Другое оборудование\n\n"
        "Просто напишите номер или название."
    )


# ============================================================
#  РАСШИРЕННЫЙ АНАЛИЗ ЗАПРОСОВ
# ============================================================

def detect_ship(text: str) -> str | None:
    """Определяет судно по тексту."""
    text_lower = text.lower()
    ships = load_ships()
    for key, name in ships.items():
        if key in text_lower:
            return name
    return None


def detect_pump_type(text: str) -> str | None:
    """Определяет тип насоса по тексту."""
    text_lower = text.lower()

    piston_keywords = ["поршн", "плунж", "прямодейств", "паровой"]
    for kw in piston_keywords:
        if kw in text_lower:
            return "piston"

    gear_keywords = ["шестерен", "шестерён", "шестерн", "ротан", "rotan", "зубчат", "маслян"]
    for kw in gear_keywords:
        if kw in text_lower:
            return "gear"

    centrifugal_keywords = ["центробеж", "центр", "крыльчатк"]
    for kw in centrifugal_keywords:
        if kw in text_lower:
            return "centrifugal"

    if "насос" in text_lower:
        if any(kw in text_lower for kw in ["маслян", "ротан"]):
            return "gear"
        if any(kw in text_lower for kw in ["крыльчатк", "центр"]):
            return "centrifugal"
        if any(kw in text_lower for kw in ["поршн", "плунж"]):
            return "piston"

    return None


def extract_equipment(text: str) -> str | None:
    """Извлекает упоминание оборудования из текста."""
    text_lower = text.lower()
    equipment_keywords = [
        "насос", "двигатель", "компрессор", "вентилятор",
        "генератор", "кран", "лебедка", "редуктор", "гидромотор",
        "брашпиль", "котёл", "водонагреватель", "дизель", "мотор",
    ]
    for kw in equipment_keywords:
        if kw in text_lower:
            pattern = r'(\w+\s+){0,2}' + kw + r'(\s+\w+){0,2}'
            match = re.search(pattern, text)
            if match:
                return match.group(0).strip()
    return None


def extract_clearances_from_text(text: str) -> list:
    """Извлекает зазоры из текста."""
    text_lower = text.lower()
    clearances = []

    clearance_map = {
        "радиальн": "radial",
        "осев": "axial",
        "подшипник": "bearing",
        "сальник": "seal",
        "цилиндр": "cylinder_piston",
        "канавк": "ring_groove",
        "замк": "ring_gap",
        "крейцкопф": "crosshead",
        "коренн": "main_bearing",
        "шатун": "connecting_rod",
        "грундбукс": "seal_wear",
    }

    patterns = [
        r'зазор\s+(\w+)\s+(\d+\.?\d*)',
        r'(\w+)\s+зазор\s+(\d+\.?\d*)',
        r'зазор\s+(\d+\.?\d*)\s+(\w+)',
        r'зазор\s+(\d+\.?\d*)',
    ]

    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            if isinstance(match, tuple):
                parts = list(match)
                value = None
                clearance_type = None
                for part in parts:
                    if re.match(r'^\d+\.?\d*$', part):
                        value = float(part)
                    else:
                        # Стем-маппинг: part может быть полным словом
                        # (например, "радиальный"), а в clearance_map — корень
                        # ("радиальн"). Сравниваем по префиксу.
                        for stem, mapped in clearance_map.items():
                            if part.startswith(stem):
                                clearance_type = mapped
                                break
                        else:
                            if part in ["radial", "axial", "bearing", "seal"]:
                                clearance_type = part
                if value is not None:
                    clearances.append({
                        "type": clearance_type or "unknown",
                        "value": value,
                        "raw": match,
                    })
            else:
                if re.match(r'^\d+\.?\d*$', match):
                    for ct, mapped in clearance_map.items():
                        if ct in text_lower:
                            clearances.append({
                                "type": mapped,
                                "value": float(match),
                                "raw": match,
                            })
                            break

    return clearances


def extract_defects(text: str) -> list:
    """Извлекает дефекты из текста."""
    text_lower = text.lower()
    defects = []

    if "дефекты" in text_lower:
        defect_part = re.split(r'дефекты[:;]', text_lower, flags=re.IGNORECASE)
        if len(defect_part) > 1:
            parts = defect_part[1].strip().split(',')
            parts = [p.strip() for p in parts if p.strip()]
            for p in parts:
                if p:
                    defects.append(p)
            if defects:
                return defects

    defect_keywords = [
        "поврежден", "повреждена", "повреждено", "повреждены",
        "сгнил", "сгнила", "сгнило", "сгнили",
        "высох", "высохла", "высохло", "высохли",
        "треснул", "треснула", "треснуло", "треснули",
        "сломан", "сломана", "сломано", "сломаны",
        "разбит", "разбита", "разбито", "разбиты",
        "износ", "течь", "коррози", "трещин", "разруш", "выкрашиван",
        "задир", "деформац", "ржав", "люфт", "биение", "стук", "вибрац",
        "зазор", "перегрев", "заедание", "отказ", "неисправн", "поломк",
        "изгиб", "скручиван", "ослаблен", "изношен", "выработк",
        "закоксовыван", "загрязнен", "неплотн", "подтекани", "протечк",
    ]

    found_defects = []
    sentences = re.split(r'[,.!?;]', text)
    for sentence in sentences:
        sentence_lower = sentence.lower().strip()
        if not sentence_lower:
            continue
        for kw in defect_keywords:
            if kw in sentence_lower:
                found_defects.append(sentence.strip())
                break

    if found_defects:
        return found_defects

    if "зазор" in text_lower:
        clearances = extract_clearances_from_text(text)
        for c in clearances:
            found_defects.append(f"зазор {c['type']}: {c['value']} мм")
        return found_defects

    return []


def parse_works_for_avr(text: str) -> list:
    """Парсит выполненные работы для АВР из текста."""
    works = []
    text_lower = text.lower()

    clean_text = re.sub(r'(авр|акт выполненных|сделай авр|создай авр|оформи авр)\s*', '', text_lower, flags=re.IGNORECASE)
    clean_text = re.sub(r'по судну\s+\w+\s*', '', clean_text)
    clean_text = re.sub(r'судно\s+\w+\s*', '', clean_text)
    clean_text = clean_text.strip()

    lines = re.split(r'\n|\.\s+|;\s+', clean_text)

    for line in lines:
        line = line.strip()
        if not line or len(line) < 5:
            continue

        work = {"name": "", "description": "", "quantity": "", "unit": "", "note": ""}

        quantity_match = re.search(r'(\d+)\s*(шт|компл|м|кг|л|шт\.|компл\.|м\.|кг\.|л\.)', line)
        if quantity_match:
            work["quantity"] = quantity_match.group(1)
            work["unit"] = quantity_match.group(2).replace('.', '')
            line = line.replace(quantity_match.group(0), '').strip()

        note_match = re.search(r'\([^)]+\)', line)
        if note_match:
            work["note"] = note_match.group(0).strip('()')
            line = line.replace(note_match.group(0), '').strip()

        if ':' in line or '—' in line or '-' in line:
            parts = re.split(r':\s*|—\s*|-\s*', line, maxsplit=1)
            if len(parts) == 2:
                work["name"] = parts[0].strip().capitalize()
                work["description"] = parts[1].strip().capitalize()
            else:
                work["description"] = line.capitalize()
        else:
            if any(word in line for word in ["замена", "ремонт", "восстановлен", "изготовлен", "монтаж", "демонтаж"]):
                work["name"] = "Ремонтные работы"
                work["description"] = line.capitalize()
            else:
                work["description"] = line.capitalize()

        if work["name"] or work["description"]:
            if not work["unit"]:
                work["unit"] = "компл." if not work["quantity"] else ""
            works.append(work)

    if not works and text.strip():
        works.append({
            "name": "Основные работы",
            "description": text.strip().capitalize(),
            "quantity": "1",
            "unit": "компл.",
            "note": "",
        })

    return works


def analyze_query_local(text: str) -> dict:
    """Локальный анализ запроса (без AI)."""
    result = {
        "ship": detect_ship(text),
        "equipment": extract_equipment(text),
        "defects": extract_defects(text),
        "pump_type": detect_pump_type(text),
        "equipment_type": detect_equipment_type(text),
        "clearances": extract_clearances_from_text(text),
        "works": parse_works_for_avr(text),
        "full_text": text,
        "source": "local",
    }

    if result["clearances"] and not result["defects"]:
        for c in result["clearances"]:
            result["defects"].append(f"зазор {c['type']}: {c['value']} мм")

    if not result["equipment"] and "насос" in text.lower():
        pump_name = pump_db.get_pump_name(result["pump_type"]) if result["pump_type"] else ""
        result["equipment"] = f"насос {pump_name}".strip() if pump_name else "насос"

    return result


# ============================================================
#  ГЕНЕРАЦИЯ ОБЪЁМА РАБОТ
# ============================================================

def generate_work_volume(defects: list, full_text: str, pump_type: str | None = None, equipment_type: str | None = None) -> str:
    """Генерирует объём работ с использованием базы знаний или AI."""
    import bot_context

    if bot_context.alisa_router:
        try:
            ai_result = bot_context.alisa_router.generate_work_volume(defects, equipment_type, pump_type)
            if ai_result:
                logger.info("Сгенерировано через Алису")
                return ai_result
        except Exception as e:
            logger.warning(f"Ошибка при вызове Алисы: {e}")

    logger.warning("Использую базовый шаблон")
    return generate_base_work_volume(defects)


def generate_base_work_volume(defects: list) -> str:
    """Базовый объём работ (запасной вариант)."""
    lines = ["1. Демонтаж узла", "2. Разборка и дефектация"]

    work_items = []
    for defect in defects:
        defect_lower = defect.lower()
        if "течь" in defect_lower or "уплотнен" in defect_lower:
            work_items.append("замена уплотнительных элементов")
        elif "износ" in defect_lower or "изношен" in defect_lower:
            work_items.append("восстановление или замена изношенных деталей")
        elif "трещин" in defect_lower:
            work_items.append("заварка трещин или замена детали")
        elif "коррози" in defect_lower:
            work_items.append("зачистка и восстановление коррозионных повреждений")
        elif "зазор" in defect_lower:
            work_items.append("регулировка зазоров")
        elif "протечк" in defect_lower:
            work_items.append("замена уплотнений и проверка герметичности")

    if work_items:
        unique_items = list(dict.fromkeys(work_items))
        lines.append("3. " + "; ".join(unique_items))
    else:
        lines.append("3. Замена/восстановление деталей")

    lines.append("4. Сборка с проверкой зазоров")
    lines.append("5. Монтаж")
    lines.append("6. Предъявление лицу сдающему")

    return "\n".join(lines)


# ============================================================
#  ПОСТРОЕНИЕ ТАБЛИЦЫ (ДЛЯ НАСОСОВ)
# ============================================================

DEFECT_MAP = {
    "крыльчатк": "2.1",
    "крылатк": "2.1",
    "колес": "2.1",
    "вал": "2.2",
    "трещин вала": "2.2",
    "шпонк": "2.3",
    "уплотнен": "3.1",
    "сальник": "3.1",
    "подшипник": "4.1",
    "крепеж": "6.3",
    "болт": "6.3",
    "гайк": "6.3",
}

DEFAULT_NO_DEFECT_TEXT = "Визуальный осмотр. Дефектов не обнаружено."
DEFAULT_NO_DEFECT_WORK = "Мыть, чистить. Годен к дальнейшей эксплуатации."

ACTION_MAP = {
    "эксплуатационный износ": "Замена.",
    "износ": "Замена.",
    "коррози": "Чистка УШМ, грунтовка. Пригодна к дальнейшей эксплуатации.",
    "грязев": "Мыть, чистить. Годен к дальнейшей эксплуатации.",
    "окисление": "Мыть, чистить. Годен к дальнейшей эксплуатации.",
    "трещин": "Замена.",
    "течь": "Замена уплотнений, проверка герметичности.",
}


def build_defect_table_pump(pump_type: str | None, defects: list, work_volume: str) -> list:
    """Строит таблицу для насосов (7 колонок)."""
    rows = [
        {"num": "1.1", "part": "Корпус насоса", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "1.2", "part": "Уплотнительное кольцо", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "1.3", "part": "Всасывающий/напорный патрубок", "defect": "", "unit": "компл.", "qty": "1"},
    ]

    if pump_type == "centrifugal":
        rows.extend([
            {"num": "2.1", "part": "Рабочее колесо (крыльчатка)", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.2", "part": "Вал насоса", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.3", "part": "Шпонка рабочего колеса", "defect": "", "unit": "шт.", "qty": "1"},
        ])
    elif pump_type == "gear":
        rows.extend([
            {"num": "2.1", "part": "Ведущая шестерня", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.2", "part": "Ведомая шестерня", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.3", "part": "Пальцы и втулки", "defect": "", "unit": "компл.", "qty": "1"},
            {"num": "2.4", "part": "Перепускной клапан", "defect": "", "unit": "шт.", "qty": "1"},
        ])
    elif pump_type == "piston":
        rows.extend([
            {"num": "2.1", "part": "Цилиндр (зеркало)", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.2", "part": "Поршень / плунжер", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.3", "part": "Поршневые кольца", "defect": "", "unit": "компл.", "qty": "1"},
            {"num": "2.4", "part": "Шток / плунжер", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.5", "part": "Крейцкопф (башмаки, направляющие)", "defect": "", "unit": "компл.", "qty": "1"},
        ])

    rows.extend([
        {"num": "3.1", "part": "Уплотнение вала (сальник/торцевое)", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "4.1", "part": "Подшипниковый узел", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "4.2", "part": "Корпус подшипников / крышки", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "4.3", "part": "Масляная камера", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "5.1", "part": "Обмотка статора", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "5.2", "part": "Клеммная коробка и кабельный ввод", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "5.3", "part": "Муфта соединения", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "6.1", "part": "Обратный клапан", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "6.2", "part": "Задвижка на всасывании", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "6.3", "part": "Крепёж и расходные материалы", "defect": "", "unit": "компл.", "qty": "1"},
    ])

    for defect in defects:
        defect_lower = defect.lower()
        placed = False
        for keyword, pos in DEFECT_MAP.items():
            if keyword in defect_lower:
                for row in rows:
                    if row["num"] == pos and not row["defect"]:
                        row["defect"] = defect
                        placed = True
                        break
                if placed:
                    break
        if not placed:
            for row in rows:
                if row["num"] == "1.1" and not row["defect"]:
                    row["defect"] = defect
                    break

    for row in rows:
        if not row["defect"]:
            row["defect"] = DEFAULT_NO_DEFECT_TEXT
            row["work"] = DEFAULT_NO_DEFECT_WORK
        else:
            defect_lower = row["defect"].lower()
            matched = None
            for keyword, action in ACTION_MAP.items():
                if keyword in defect_lower:
                    matched = action
                    break
            row["work"] = matched or work_volume

    return rows


# ============================================================
#  ПОСТРОЕНИЕ ТАБЛИЦЫ (ДЛЯ ДВИГАТЕЛЕЙ)
# ============================================================

def build_defect_table_engine(defects: list, work_volume: str) -> list:
    """Строит таблицу для двигателей (6 колонок)."""
    rows = []

    sections = {
        "Цилиндропоршневая группа": [],
        "Головка цилиндров и газораспределение": [],
        "Системы и вспомогательное оборудование": [],
        "Сборочные и испытательные работы": [],
    }

    for i, defect in enumerate(defects, 1):
        defect_lower = defect.lower()
        section = None

        if any(word in defect_lower for word in ["поршн", "цилиндр", "кольц", "втулк"]):
            section = "Цилиндропоршневая группа"
        elif any(word in defect_lower for word in ["крышк", "клапан", "форсунк", "толкател", "газораспредел"]):
            section = "Головка цилиндров и газораспределение"
        elif any(word in defect_lower for word in ["рубашк", "сальник", "турбо", "масл", "охлажд"]):
            section = "Системы и вспомогательное оборудование"
        elif any(word in defect_lower for word in ["испытани", "сборк"]):
            section = "Сборочные и испытательные работы"
        else:
            section = "Прочее"

        if not defect or defect == "Не указано":
            defect_text = DEFAULT_NO_DEFECT_TEXT
            work_text = DEFAULT_NO_DEFECT_WORK
        else:
            defect_lower = defect.lower()
            matched = None
            for keyword, action in ACTION_MAP.items():
                if keyword in defect_lower:
                    matched = action
                    break
            defect_text = defect
            work_text = matched or work_volume

        rows.append({
            "num": str(i),
            "defect": defect_text,
            "work": work_text,
            "unit": "компл.",
            "qty": "1",
            "section": section,
        })

    return rows


# ============================================================
#  ФУНКЦИИ ORM: НАГРУЗКА РЕМОНТНОЙ ВЕДОМОСТИ
# ============================================================

def save_repair_items_to_db(ship_id: int, items: list) -> tuple:
    """
    Сохраняет пункты ремонтной ведомости в БД с дедупликацией.
    Возвращает: (inserted_count, skipped_count, statement_id)
    """
    session = SessionLocal()
    try:
        stmt = RepairStatement(ship_id=ship_id, source_excel_file_ref="uploaded")
        session.add(stmt)
        session.flush()
        statement_id = stmt.id

        inserted = 0
        skipped = 0

        for item in items:
            existing = session.query(StatementItem).filter(
                StatementItem.statement_id == statement_id,
                StatementItem.item_number == item.get("item_number"),
                StatementItem.section == item.get("section"),
            ).first()

            if existing:
                skipped += 1
                continue

            stmt_item = StatementItem(
                statement_id=statement_id,
                item_number=item.get("item_number"),
                description=item.get("description"),
                quantity=item.get("quantity"),
                section=item.get("section"),
                status="active",
            )
            session.add(stmt_item)
            inserted += 1

        session.commit()
        return inserted, skipped, statement_id
    except Exception as e:
        session.rollback()
        raise e
    finally:
        session.close()


def get_user_role(telegram_id: int) -> str:
    """Получить роль пользователя из ORM."""
    session = SessionLocal()
    try:
        user = session.query(User).filter_by(telegram_id=telegram_id).first()
        return user.role if user else "customer"
    finally:
        session.close()


def can_upload_repair_list(telegram_id: int) -> bool:
    """Проверить, может ли пользователь загружать ремонтную ведомость.
    Доступно всем, кроме customer.
    """
    role = get_user_role(telegram_id)
    return role != "customer"


# ============================================================
#  ФУНКЦИИ ВЕРСИОНИРОВАНИЯ ДОКУМЕНТОВ
# ============================================================

def handle_document_approve(document_id: int, user_id: int) -> tuple:
    """
    Утвердить документ: draft → approved.
    Возвращает: (success, message)
    """
    return approve_document(document_id, user_id)


def handle_document_archive(document_id: int, user_id: int, admin_ids: list | None = None) -> tuple:
    """
    Архивировать документ: approved → archived.
    Только ADMIN_IDS.
    Возвращает: (success, message)
    """
    return archive_document(document_id, user_id, admin_ids)


def handle_document_delete(document_id: int, user_id: int, admin_ids: list | None = None) -> tuple:
    """
    Удалить документ.
    - draft: любой может удалить
    - approved: только ADMIN_IDS
    Возвращает: (success, message)
    """
    return delete_document(document_id, user_id, admin_ids)


# Экземпляр базы данных насосов (для обратной совместимости)
pump_db = PumpDatabase()
