# -*- coding: utf-8 -*-
"""
Создание документов Word (акты дефектации и акты выполненных работ).

Содержит функции построения таблиц дефектации и генерации .docx-файлов
на основе шаблонов. Импортирует детекцию, справочники и замену плейсхолдеров
из services.extra, а нумерацию — из отдельного сервиса счётчиков.
"""

from io import BytesIO
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from services.extra import (
    detect_equipment_type,
    build_defect_table_pump,
    build_defect_table_engine,
)
from services.catalog_service import load_companies
from services.document_counter_service import get_next_number
from services.template_service import load_template, replace_placeholders


def create_defect_document(
    ship: str,
    equipment: str,
    defects: list,
    work_volume: str,
    pump_type: str | None = None,
    repair_type: str | None = None,
    purpose: str | None = None,
    basis: str | None = None,
) -> BytesIO:
    """Создаёт акт дефектации с таблицей, подходящей под тип оборудования.

    Args:
        ship: Название судна.
        equipment: Название оборудования.
        defects: Список дефектов.
        work_volume: Объём работ.
        pump_type: Тип насоса (для насосов).
        repair_type: Тип ремонта.
        purpose: Цель акта.
        basis: Основание для акта.

    Returns:
        BytesIO с содержимым .docx-файла.
    """
    doc = load_template("defect_act_template.docx")

    number = get_next_number("da")

    date_str = datetime.now().strftime('%d.%m.%Y')
    ship_code = ship[:3].upper() if ship else "XXX"

    equipment_type = detect_equipment_type(equipment or "")
    if equipment_type is None:
        equipment_type = "pump"

    if equipment_type == "pump":
        rows_data = build_defect_table_pump(pump_type, defects, work_volume)
        cols = 7
        headers = ['№', 'Позиция', 'Дефект / Состояние', 'Объём работ', 'Ед. изм.', 'Кол-во', 'Примечание']
        sections = {
            "1": "Корпус и проточная часть",
            "2": "Ротор / рабочая часть",
            "3": "Уплотнения вала",
            "4": "Подшипниковый узел",
            "5": "Электропривод",
            "6": "Арматура и обвязка",
        }
        get_section_key = lambda row: row["num"].split(".")[0]
        show_purpose = True
        show_basis = True
        show_conclusion = True
        show_notes = False
        notes_text = ""
    else:
        rows_data = build_defect_table_engine(defects, work_volume)
        cols = 6
        headers = ['№ п/п', 'Наименование дефекта', 'Объём работ', 'Ед. изм.', 'Кол-во', 'Примечание']
        sections = {}
        get_section_key = lambda row: row.get("section", "Прочее")
        show_purpose = False
        show_basis = False
        show_conclusion = False
        show_notes = True
        notes_text = (
            "Все СЗЧ (поршневые кольца, поршни, втулка, комплекты для форсунок, РТИ) — "
            "поставка Заказчика, если не указано иное.\n"
            "Работы по проточке и транспортировке деталей выполняются Подрядчиком "
            "за отдельную плату (акт дополнительных работ)."
        )

    table_paragraph_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "{{table}}" in paragraph.text:
            table_paragraph_index = i
            paragraph.text = ""
            break

    table = doc.add_table(rows=1, cols=cols)
    table.autofit = False
    table.allow_autofit = False

    if cols == 7:
        widths = [Cm(1.3), Cm(3.8), Cm(5.0), Cm(5.0), Cm(2.0), Cm(1.8), Cm(3.8)]
    else:
        widths = [Cm(1.8), Cm(5.0), Cm(6.0), Cm(2.2), Cm(2.0), Cm(3.0)]

    for i, width in enumerate(widths):
        table.columns[i].width = width

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    current_section = None
    for row_data in rows_data:
        section_key = get_section_key(row_data)
        if section_key != current_section:
            current_section = section_key
            row = table.add_row().cells
            for cell in row:
                cell.text = ""
            if cols == 7:
                row[0].text = sections.get(section_key, "")
            else:
                row[0].text = section_key
            for cell in row:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        row = table.add_row().cells
        if cols == 7:
            row[0].text = row_data["num"]
            row[1].text = row_data["part"]
            row[2].text = row_data.get("defect", "—")
            row[3].text = row_data.get("work", "—")
            row[4].text = row_data["unit"]
            row[5].text = row_data["qty"]
            row[6].text = "—"
        else:
            row[0].text = row_data["num"]
            row[1].text = row_data.get("defect", "—")
            row[2].text = row_data.get("work", "—")
            row[3].text = row_data.get("unit", "компл.")
            row[4].text = row_data.get("qty", "1")
            row[5].text = "—"

    if table_paragraph_index is not None:
        target_paragraph = doc.paragraphs[table_paragraph_index]
        tbl = table._tbl
        target_paragraph._element.addprevious(tbl)

    placeholders = {
        "ship_code": ship_code,
        "number": str(number).zfill(2),
        "date": date_str,
        "ship": ship or "Не указано",
        "equipment": equipment or "Не указано",
        "work_object": repair_type or "Текущий ремонт",
        "purpose": purpose or "Определение технического состояния и объема ремонта",
        "basis": basis or f"План-график ремонта на {datetime.now().year} год",
    }

    if show_conclusion:
        placeholders["conclusion"] = "Детали подлежат замене/восстановлению согласно указанному объёму работ."
    placeholders["special_notes"] = notes_text if show_notes else ""

    doc = replace_placeholders(doc, placeholders)

    if equipment_type != "pump":
        for paragraph in doc.paragraphs:
            if "Представитель подрядчика (Исполнитель)" in paragraph.text:
                paragraph.text = "Представитель Подрядчика:"
            if "Представитель заказчика (Судовладелец / Экипаж)" in paragraph.text:
                paragraph.text = "Представитель Заказчика:"
            if "Старший механик" in paragraph.text:
                paragraph.text = "Должность      / *[ФИО]* /"
            if "Согласовано (при необходимости)" in paragraph.text:
                paragraph.text = ""
            if "Инспектор РМРС" in paragraph.text:
                paragraph.text = ""

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def create_avr_document(
    ship: str,
    works: list,
    executor: str | None = None,
    customer: str | None = None,
    location: str | None = None,
) -> BytesIO:
    """Создаёт акт выполненных работ (АВР).

    Args:
        ship: Название судна.
        works: Список выполненных работ.
        executor: Исполнитель (по умолчанию из companies.json).
        customer: Заказчик (по умолчанию из companies.json).
        location: Место проведения работ (по умолчанию из companies.json).

    Returns:
        BytesIO с содержимым .docx-файла.
    """
    companies = load_companies()
    executor = executor or companies.get("executor")
    customer = customer or companies.get("customer")
    location = location or companies.get("location")
    doc = load_template("avr_template.docx")

    number = get_next_number("avr")

    date_str = datetime.now().strftime('%d.%m.%Y')
    ship_code = ship[:3].upper() if ship else "XXX"

    table_paragraph_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "{{table}}" in paragraph.text:
            table_paragraph_index = i
            paragraph.text = ""
            break

    table = doc.add_table(rows=1, cols=6)
    table.autofit = False
    table.allow_autofit = False

    widths = [Cm(1.8), Cm(5.0), Cm(6.0), Cm(2.2), Cm(2.0), Cm(3.0)]
    for i, width in enumerate(widths):
        table.columns[i].width = width

    headers = ['№ п/п', 'Наименование работ', 'Описание выполненных работ', 'Кол-во', 'Ед. изм.', 'Примечание']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if works:
        for i, work in enumerate(works, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = work.get('name', '')
            row[2].text = work.get('description', '')
            row[3].text = str(work.get('quantity', ''))
            row[4].text = work.get('unit', '')
            row[5].text = work.get('note', '')
    else:
        row = table.add_row().cells
        row[0].text = "1"
        row[1].text = "Основные работы"
        row[2].text = "Выполнены работы согласно дефектации"
        row[3].text = "1"
        row[4].text = "компл."
        row[5].text = ""

    if table_paragraph_index is not None:
        target_paragraph = doc.paragraphs[table_paragraph_index]
        tbl = table._tbl
        target_paragraph._element.addprevious(tbl)

    placeholders = {
        "ship_code": ship_code,
        "number": str(number).zfill(2),
        "date": date_str,
        "ship": ship or "Не указано",
        "executor": executor,
        "customer": customer,
        "location": location,
    }

    doc = replace_placeholders(doc, placeholders)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
