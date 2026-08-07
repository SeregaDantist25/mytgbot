# scripts/convert_word_to_json.py
import os
import json
import re
from pathlib import Path
from docx import Document

# Ключевые слова для распознавания заголовка таблицы.
# Реальный шаблон компании использует "№" + "Позиция", а не "№ п/п" + "Наименование" —
# поэтому проверяем набор характерных слов таблицы дефектации, а не одну жёсткую фразу.
HEADER_KEYWORDS = ["позиция", "наименование", "узел", "детал"]
DEFECT_HEADER_KEYWORDS = ["дефект", "состояние", "описание"]
WORK_HEADER_KEYWORDS = ["объём работ", "объем работ", "ремонтных работ", "наименование работ"]


def _is_table_header_row(row):
    """Строка считается заголовком таблицы, только если в ней одновременно
    встречается колонка-позиция И колонка-дефект/работы. Одного слова недостаточно —
    иначе случайно словим обычную строку данных."""
    row_text = " ".join(row).lower()
    has_position_col = any(kw in row_text for kw in HEADER_KEYWORDS) or row[0].strip() == "№"
    has_defect_or_work_col = any(kw in row_text for kw in DEFECT_HEADER_KEYWORDS) or \
        any(kw in row_text for kw in WORK_HEADER_KEYWORDS)
    return has_position_col and has_defect_or_work_col


def _map_columns(headers):
    """headers — список строк заголовка таблицы (уже lower())."""
    col_map = {}
    for i, header in enumerate(headers):
        if i == 0 and header.strip() in ("№", "№ п/п"):
            continue
        if any(kw in header for kw in ["позиция", "наименование", "узел", "детал"]):
            col_map.setdefault("part", i)
        elif any(kw in header for kw in ["дефект", "описание", "состояние"]):
            col_map.setdefault("defect", i)
        elif any(kw in header for kw in ["объём работ", "объем работ", "ремонт", "наименование работ"]):
            col_map.setdefault("work", i)
        elif "ед" in header:
            col_map.setdefault("unit", i)
        elif "кол" in header:
            col_map.setdefault("qty", i)
        elif "примечан" in header:
            col_map.setdefault("note", i)
    return col_map


def parse_word_act(file_path):
    """Парсит Word-файл с актом дефектации"""
    try:
        doc = Document(file_path)

        full_text = []
        tables_data = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())

        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                # ВАЖНО: раньше пустые ячейки просто выбрасывались (row_text.append только
                # непустых значений), из-за чего вся строка "съезжала" по колонкам.
                # Сохраняем позиционное соответствие — пустая ячейка остаётся пустой строкой.
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    table_rows.append(row_text)
            if table_rows:
                tables_data.append(table_rows)

        full_text_str = "\n".join(full_text)

        ship = "Не указано"
        equipment = "Не указано"
        repair_type = "Текущий ремонт"
        rows_data = []

        ship_patterns = [
            r'Судно[:\s]+([^\n]+)',
            r'(?:т/х|бк)\s*[«"]([^»"]+)[»"]',
        ]
        for pattern in ship_patterns:
            match = re.search(pattern, full_text_str, re.IGNORECASE)
            if match:
                ship = match.group(1).strip()
                break

        equip_patterns = [
            r'Оборудование[:\s]+([^\n]+)',
            r'Механизм[:\s]+([^\n]+)',
        ]
        for pattern in equip_patterns:
            match = re.search(pattern, full_text_str, re.IGNORECASE)
            if match:
                equipment = match.group(1).strip()
                break

        repair_match = re.search(r'Категория\s*ремонта[:\s]+([^\n]+)', full_text_str, re.IGNORECASE)
        if repair_match:
            repair_type = repair_match.group(1).strip()

        for table in tables_data:
            header_row_idx = None
            for i, row in enumerate(table):
                if _is_table_header_row(row):
                    header_row_idx = i
                    break

            if header_row_idx is None:
                continue

            headers = [h.strip().lower() for h in table[header_row_idx]]
            col_map = _map_columns(headers)

            if "part" not in col_map:
                col_map["part"] = 1 if len(headers) > 1 else 0

            for row in table[header_row_idx + 1:]:
                def cell(key, default=""):
                    idx = col_map.get(key)
                    if idx is None or idx >= len(row):
                        return default
                    return row[idx].strip()

                part = cell("part")
                if not part or len(part) < 2:
                    continue
                if re.match(r'^\d+[.,]?\d*$', part):
                    continue

                defect = cell("defect")
                work = cell("work")
                unit = cell("unit")
                qty = cell("qty")
                note = cell("note")

                rows_data.append({
                    "part": part,
                    "defect": defect if defect and defect != "—" else "Визуальный осмотр. Дефектов не обнаружено.",
                    "work": work if work and work != "—" else "",
                    "unit": unit,
                    "qty": qty,
                    "note": note,
                })

        # Схема "rows" — новая, структурированная (по колонкам).
        # Схема "defects"/"work_volume" — оставлена для обратной совместимости с
        # knowledge_from_acts.py и ai_router.py, которые пока читают старый формат.
        act_data = {
            "ship": ship,
            "equipment": equipment,
            "repair_type": repair_type,
            "rows": rows_data,
            "defects": [f"{r['part']}: {r['defect']}" for r in rows_data],
            "work_volume": "\n".join(f"{r['part']}: {r['work']}" for r in rows_data if r["work"]),
            "conclusion": "Детали подлежат замене/восстановлению согласно указанному объёму работ.",
        }

        if rows_data:
            return [act_data]

        print(f"  ⚠️ Таблица не распознана в {file_path.name} — проверьте формат вручную")
        return []

    except Exception as e:
        print(f"❌ Ошибка {file_path.name}: {e}")
        return []


def convert_all_word():
    input_dir = Path("data/acts_word")
    output_dir = Path("data/act_examples")
    output_dir.mkdir(parents=True, exist_ok=True)

    if not input_dir.exists():
        print(f"⚠️ Папка {input_dir} не найдена. Создайте её и положите туда Word-файлы.")
        return

    total = 0
    for file in input_dir.glob("*.docx"):
        print(f"📄 {file.name}")
        results = parse_word_act(file)
        if results:
            for i, act in enumerate(results, 1):
                out_file = output_dir / f"{file.stem}_{i}.json"
                with open(out_file, 'w', encoding='utf-8') as f:
                    json.dump(act, f, ensure_ascii=False, indent=2)
                total += 1
            print(f"  ✅ {len(results)} актов, {len(results[0]['rows'])} позиций")
        else:
            print(f"  ⚠️ Нет данных")

    print(f"\n🎉 Всего сохранено: {total} актов")


if __name__ == "__main__":
    convert_all_word()
