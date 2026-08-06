# models/alisa_act_creator.py
import os
import json
import httpx
import re
from typing import Dict, List, Optional
from docx import Document

class AlisaActCreator:
    def __init__(self):
        self.api_key = os.environ.get('YANDEX_API_KEY')
        self.folder_id = os.environ.get('YANDEX_FOLDER_ID')
        self.templates_dir = "templates"
        self.examples_dir = "data/act_examples"
        print(f"🔧 AlisaActCreator инициализирован. API Key: {'есть' if self.api_key else 'НЕТ'}")
    
    def _load_template_text(self) -> str:
        """Загружает текст шаблона для контекста Алисы"""
        try:
            doc = Document(os.path.join(self.templates_dir, "defect_act_template.docx"))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            if table_text:
                full_text += "\n\nСТРУКТУРА ТАБЛИЦЫ:\n" + "\n".join(table_text)
            
            return full_text
        except Exception as e:
            print(f"⚠️ Ошибка загрузки шаблона: {e}")
            return "Шаблон не загружен"
    
    def _load_examples(self) -> str:
        """Загружает примеры актов"""
        examples = []
        try:
            if os.path.exists(self.examples_dir):
                for file in os.listdir(self.examples_dir):
                    if file.endswith(".json"):
                        with open(os.path.join(self.examples_dir, file), 'r', encoding='utf-8') as f:
                            data = json.load(f)
                            examples.append(json.dumps(data, ensure_ascii=False, indent=2))
            return "\n\n---\n\n".join(examples[:5]) if examples else "Примеры не загружены"
        except Exception as e:
            print(f"⚠️ Ошибка загрузки примеров: {e}")
            return "Примеры не загружены"
    
    def generate_act_data(self, user_text: str) -> Dict:
        """Генерация данных для акта через Алису"""
        try:
            template_text = self._load_template_text()
            examples_text = self._load_examples()
            
            prompt = f"""Ты — инженерный ассистент для судоремонта. Твоя задача — проанализировать запрос пользователя и сгенерировать данные для Акта дефектации.

## ШАБЛОН АКТА ДЕФЕКТАЦИИ:

{template_text[:2000]}

## ПРИМЕРЫ ЗАПОЛНЕННЫХ АКТОВ:

{examples_text[:2000]}

## ЗАПРОС ПОЛЬЗОВАТЕЛЯ:

{user_text}

## ИНСТРУКЦИЯ:

1. Извлеки из запроса:
   - Название судна (ship)
   - Оборудование (equipment) — тип и название
   - Дефекты (defects) — список
   - Тип ремонта (repair_type) — если указан

2. Сгенерируй объём работ (work_volume) — нумерованный список

3. ОТВЕТЬ ТОЛЬКО JSON-ОБЪЕКТОМ:

{{
  "ship": "название_судна",
  "equipment": "название_оборудования",
  "repair_type": "текущий/капитальный/средний",
  "defects": ["дефект1", "дефект2"],
  "work_volume": "1. Демонтаж...\\n2. Разборка...\\n3. ...",
  "conclusion": "Детали подлежат замене/восстановлению..."
}}

ОТВЕТЬ ТОЛЬКО JSON-ОБЪЕКТОМ, БЕЗ ЛИШНЕГО ТЕКСТА!"""

            print(f"🧠 Отправляем запрос в Алису для генерации акта...")
            response = self._call_yandex_gpt(prompt)
            
            if not response:
                print(f"⚠️ Алиса не ответила, использую локальный парсер")
                return self._fallback_act_data(user_text)
            
            return self._parse_act_data(response, user_text)
            
        except Exception as e:
            print(f"⚠️ Ошибка генерации акта через Алису: {e}")
            return self._fallback_act_data(user_text)
    
    def _call_yandex_gpt(self, prompt: str) -> Optional[str]:
        """Вызов YandexGPT"""
        try:
            with httpx.Client(timeout=45.0) as client:
                response = client.post(
                    "https://llm.api.cloud.yandex.net/foundationModels/v1/completion",
                    headers={
                        "Authorization": f"Api-Key {self.api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "modelUri": f"gpt://{self.folder_id}/yandexgpt/latest",
                        "completionOptions": {
                            "stream": False,
                            "temperature": 0.3,
                            "maxTokens": 2000
                        },
                        "messages": [
                            {
                                "role": "system",
                                "text": "Ты — помощник для судоремонта. Генерируй только JSON-ответы."
                            },
                            {
                                "role": "user",
                                "text": prompt
                            }
                        ]
                    }
                )
                
                if response.status_code == 200:
                    result = response.json()
                    return result.get("result", {}).get("alternatives", [{}])[0].get("message", {}).get("text", "")
                else:
                    print(f"⚠️ Ошибка YandexGPT: {response.status_code}")
                    return None
                    
        except Exception as e:
            print(f"⚠️ Ошибка вызова YandexGPT: {e}")
            return None
    
    def _parse_act_data(self, response: str, fallback_text: str) -> Dict:
        """Парсинг JSON из ответа Алисы"""
        try:
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                data = json.loads(json_match.group())
                if "ship" in data and "equipment" in data and "defects" in data:
                    return data
        except Exception as e:
            print(f"⚠️ Ошибка парсинга JSON: {e}")
        
        return self._fallback_act_data(fallback_text)
    
    def _fallback_act_data(self, text: str) -> Dict:
        """Запасной вариант — локальный парсер (без циклического импорта)"""
        # Локальные функции, чтобы не импортировать из bot.py
        import re
        
        # Определяем судно
        ships = ["аргака", "пластун", "славянская", "первоуральск", "керчь", "краснодар", "алеут"]
        ship = "Не указано"
        for s in ships:
            if s in text.lower():
                ship = s.capitalize()
                break
        
        # Определяем оборудование
        equipment_keywords = ["насос", "двигатель", "компрессор", "вентилятор", "генератор"]
        equipment = "Не указано"
        for kw in equipment_keywords:
            if kw in text.lower():
                equipment = kw.capitalize()
                break
        
        # Извлекаем дефекты
        defects = []
        defect_keywords = ["износ", "течь", "коррози", "трещин", "зазор", "люфт", "биение"]
        sentences = re.split(r'[,.!?;]', text)
        for sentence in sentences:
            for kw in defect_keywords:
                if kw in sentence.lower():
                    defects.append(sentence.strip())
                    break
        
        if not defects:
            defects = ["Не указано"]
        
        # Генерируем объём работ
        work_lines = ["1. Демонтаж узла", "2. Разборка и дефектация"]
        for defect in defects:
            if "износ" in defect.lower():
                work_lines.append("3. Замена/восстановление изношенных деталей")
                break
        else:
            work_lines.append("3. Замена/восстановление деталей")
        work_lines.append("4. Сборка с проверкой зазоров")
        work_lines.append("5. Монтаж")
        work_lines.append("6. Предъявление лицу сдающему")
        
        return {
            "ship": ship,
            "equipment": equipment if equipment != "Не указано" else "Оборудование",
            "repair_type": "Текущий ремонт",
            "defects": defects,
            "work_volume": "\n".join(work_lines),
            "conclusion": "Детали подлежат замене/восстановлению согласно указанному объёму работ."
        }

# Создаём глобальный экземпляр
act_creator = AlisaActCreator()