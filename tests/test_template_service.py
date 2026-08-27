# -*- coding: utf-8 -*-

from docx import Document

from services.template_service import TemplateService


def test_load_template_from_configured_directory(tmp_path):
    source = tmp_path / "template.docx"
    Document().save(source)

    loaded = TemplateService(tmp_path).load("template.docx")
    assert loaded is not None


def test_missing_template_has_clear_error(tmp_path):
    service = TemplateService(tmp_path)
    try:
        service.load("missing.docx")
    except FileNotFoundError as error:
        assert "missing.docx" in str(error)
        assert str(tmp_path) in str(error)
    else:
        raise AssertionError("Ожидалась ошибка отсутствующего шаблона")


def test_replaces_split_placeholder_in_paragraph():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Судно: {{sh")
    paragraph.add_run("ip}}")

    TemplateService(".").replace_placeholders(document, {"ship": "Аргака"})
    assert paragraph.text == "Судно: Аргака"


def test_replaces_placeholder_inside_table():
    document = Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "Заказ: {{order}}"

    TemplateService(".").replace_placeholders(document, {"order": "24-01"})
    assert cell.text == "Заказ: 24-01"


def test_legacy_extra_imports_point_to_template_service():
    from services import extra, template_service

    assert extra.load_template is template_service.load_template
    assert extra.replace_placeholders is template_service.replace_placeholders
