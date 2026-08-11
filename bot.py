import os
import telebot
from telebot.handler_backends import State, StatesGroup
from telebot import custom_filters
import httpx
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import re
import json
import time
import sqlite3
import threading
import db
from models import init_models, sync_ships_from_json, SessionLocal, User, Ship, RepairStatement, StatementItem, Document
from file_storage import storage
import navigation
import document_commands

# Новые модули для документооборота
try:
    import document_manager
    import bot_handlers_new
    DOCUMENT_MANAGER_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Модули документооборота не загружены: {e}")
    DOCUMENT_MANAGER_AVAILABLE = False

# scanner импортируется лениво (внутри функций), т.к. требует openpyxl,
# который может отсутствовать на сервере при старте.

# --- Настройки ---
BOT_TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN')
GROQ_API_KEY = os.environ.get('GROQ_API_KEY')
# Секретный код инженера-технолога (задаётся через переменную окружения ENGINEER_CODE)
ENGINEER_CODE = os.environ.get('ENGINEER_CODE')

# ID администраторов (через запятую), имеющих доступ к /set_role
ADMIN_IDS = [int(x.strip()) for x in os.environ.get('ADMIN_IDS', '').split(',') if x.strip()]
if not BOT_TOKEN:
    raise ValueError("Переменная TELEGRAM_BOT_TOKEN не найдена!")

bot = telebot.TeleBot(BOT_TOKEN)
bot.add_custom_filter(custom_filters.StateFilter(bot))

# --- Пути к файлам ---
TEMPLATES_DIR = "templates"
DATA_DIR = os.getenv("DATA_DIR", "data")
CHECKLISTS_FILE = os.path.join(DATA_DIR, "checklists.json")
COUNTERS_FILE = os.path.join(DATA_DIR, "counters.json")
SHIPS_FILE = os.path.join(DATA_DIR, "ships.json")
COMPANIES_FILE = os.path.join(DATA_DIR, "companies.json")
COUNTERS_DB = os.path.join(DATA_DIR, "counters.db")
CHAT_STATE_FILE = os.path.join(DATA_DIR, "chat_state.json")

# ============================================================
#  ИМПОРТ ГОСТ ЧЕКЕРА И АЛИСЫ
# ============================================================

try:
    from gost_checker import GOSTChecker
    gost_checker = GOSTChecker()
    print(f"✅ ГОСТ чекер загружен. Доступно ГОСТов: {len(gost_checker.get_all_gosts())}")
except Exception as e:
    print(f"⚠️ Ошибка при загрузке ГОСТ чекера: {e}")
    gost_checker = None

# Пытаемся загрузить Алису (ai_router)
alisa_router = None
try:
    from models.ai_router import router as alisa_router
    print(f"✅ Алиса (YandexGPT) загружена успешно!")
except ImportError as e:
    print(f"⚠️ Модуль ai_router не найден: {e}")
except Exception as e:
    print(f"⚠️ Ошибка при загрузке Алисы: {e}")

# Пытаемся загрузить создателя актов через Алису
alisa_act_creator = None
try:
    from models.alisa_act_creator import act_creator as alisa_act_creator
    print(f"✅ Создатель актов через Алису загружен!")
except ImportError as e:
    print(f"⚠️ Модуль alisa_act_creator не найден: {e}")
except Exception as e:
    print(f"⚠️ Ошибка при загрузке создателя актов: {e}")

# ============================================================
#  ЗАГРУЗКА ДАННЫХ ИЗ JSON
# ============================================================

def load_checklists():
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    
    if not os.path.exists(CHECKLISTS_FILE):
        raise FileNotFoundError(f"Файл {CHECKLISTS_FILE} не найден!")
    
    with open(CHECKLISTS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_ships():
    """Загружает словарь судов из data/ships.json"""
    if not os.path.exists(SHIPS_FILE):
        return {}
    with open(SHIPS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_companies():
    """Загружает дефолтные executor/customer/location из data/companies.json"""
    defaults = {
        "executor": "ООО «Новое время»",
        "customer": "АО «Бункерная компания»",
        "location": "Рейд 4ый район, г. Находка",
    }
    if not os.path.exists(COMPANIES_FILE):
        return defaults
    with open(COMPANIES_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    defaults.update(data)
    return defaults

class PumpDatabase:
    def __init__(self):
        self.data = load_checklists()
    
    def get_pump_types(self):
        return list(self.data.keys())
    
    def get_pump_name(self, pump_type):
        return self.data.get(pump_type, {}).get("name", pump_type)
    
    def get_checklist(self, pump_type):
        return self.data.get(pump_type, {}).get("items", [])
    
    def get_clearances(self, pump_type, clearance_type):
        clearances = self.data.get(pump_type, {}).get("clearances", {})
        return clearances.get(clearance_type)
    
    def check_clearance(self, pump_type, clearance_type, measured_value):
        clearance_data = self.get_clearances(pump_type, clearance_type)
        if not clearance_data:
            return {
                "status": "unknown",
                "message": f"Данные по зазору '{clearance_type}' для '{pump_type}' отсутствуют",
                "action": "Проверьте правильность ввода"
            }
        
        standard_min = clearance_data.get("min", 0)
        standard_max = clearance_data.get("max", 0)
        unit = clearance_data.get("unit", "мм")
        
        if "мм/мм" in unit:
            return {
                "status": "info",
                "message": f"📌 Зазор зависит от диаметра: {standard_min}-{standard_max} {unit}",
                "action": "Уточните диаметр для точного расчёта"
            }
        
        if measured_value < standard_min:
            return {
                "status": "warning",
                "message": f"⚠️ Зазор МЕНЬШЕ нормы: {measured_value} мм (норма: {standard_min}-{standard_max} мм)",
                "action": "Проверьте точность измерения"
            }
        elif measured_value <= standard_max:
            return {
                "status": "ok",
                "message": f"✅ Зазор В НОРМЕ: {measured_value} мм (норма: {standard_min}-{standard_max} мм)",
                "action": "Деталь работоспособна"
            }
        else:
            return {
                "status": "critical",
                "message": f"🔴 Зазор ПРЕВЫШЕН: {measured_value} мм (норма: {standard_min}-{standard_max} мм)",
                "action": "Требуется ремонт"
            }
    
    def get_common_defects(self, pump_type):
        return self.data.get(pump_type, {}).get("defects", [])
    
    def get_repair_method(self, pump_type, defect_text):
        defect_lower = defect_text.lower()
        methods = self.data.get(pump_type, {}).get("repair_methods", {})
        for key, method in methods.items():
            if key in defect_lower:
                return method
        return None

pump_db = PumpDatabase()

# ============================================================
#  РАБОТА С ШАБЛОНАМИ
# ============================================================

def load_template(filename):
    template_path = os.path.join(TEMPLATES_DIR, filename)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон {filename} не найден в {TEMPLATES_DIR}")
    return Document(template_path)

def _merge_runs_with_tag(paragraph, tag):
    """Склеивает runs параграфа в один, если тег разбит на несколько runs."""
    full_text = paragraph.text
    if tag not in full_text:
        return
    # Проверяем, есть ли тег хотя бы в одном run целиком
    for run in paragraph.runs:
        if tag in run.text:
            return  # тег уже в одном run — нормализация не нужна
    # Тег разбит — объединяем все runs в первый, остальные обнуляем
    if not paragraph.runs:
        return
    first_run = paragraph.runs[0]
    first_run.text = full_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_placeholders(doc, placeholders):
    def _replace_in_paragraph(paragraph):
        for key in placeholders:
            tag = f"{{{{{key}}}}}"
            if tag in paragraph.text:
                _merge_runs_with_tag(paragraph, tag)
                for run in paragraph.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, str(placeholders[key]))

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph)

    return doc

def _init_counters_db():
    """Создаёт таблицу counters и мигрирует значения из counters.json один раз."""
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    conn = sqlite3.connect(COUNTERS_DB)
    cursor = conn.cursor()
    cursor.execute(
        "CREATE TABLE IF NOT EXISTS counters (doc_type TEXT PRIMARY KEY, value INTEGER)"
    )
    # Миграция из counters.json (один раз)
    if os.path.exists(COUNTERS_FILE):
        try:
            with open(COUNTERS_FILE, 'r', encoding='utf-8') as f:
                old = json.load(f)
            for doc_type, value in old.items():
                cursor.execute(
                    "INSERT OR IGNORE INTO counters (doc_type, value) VALUES (?, ?)",
                    (doc_type, value)
                )
            conn.commit()
            # После миграции переименовываем json, чтобы не мигрировать повторно
            os.rename(COUNTERS_FILE, COUNTERS_FILE + ".migrated")
        except Exception as e:
            print(f"⚠️ Ошибка миграции счётчиков: {e}")
    conn.commit()
    conn.close()

_init_counters_db()

def get_next_number(doc_type):
    """Атомарно инкрементирует счётчик и возвращает новое значение."""
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    conn = sqlite3.connect(COUNTERS_DB)
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE counters SET value = value + 1 WHERE doc_type = ? RETURNING value",
        (doc_type,)
    )
    row = cursor.fetchone()
    if row is None:
        cursor.execute("INSERT INTO counters (doc_type, value) VALUES (?, 1)", (doc_type,))
        conn.commit()
        conn.close()
        return 1
    conn.commit()
    conn.close()
    return row[0]

def get_counter(doc_type):
    """Обратная совместимость: возвращает следующий номер без инкремента."""
    return get_next_number(doc_type)

def update_counter(doc_type, new_number):
    """Обратная совместимость: устанавливает счётчик в заданное значение."""
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    conn = sqlite3.connect(COUNTERS_DB)
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO counters (doc_type, value) VALUES (?, ?) "
        "ON CONFLICT(doc_type) DO UPDATE SET value = excluded.value",
        (doc_type, new_number)
    )
    conn.commit()
    conn.close()

# ============================================================
#  ПЕРСИСТЕНТНОЕ ХРАНИЛИЩЕ СОСТОЯНИЯ ДИАЛОГА
# ============================================================

_chat_state_lock = threading.Lock()

def _load_chat_state():
    """Загружает chat_state.json в память (с блокировкой)."""
    with _chat_state_lock:
        if not os.path.exists(CHAT_STATE_FILE):
            return {}
        try:
            with open(CHAT_STATE_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}

def _save_chat_state(state):
    """Сохраняет chat_state.json на диск (с блокировкой)."""
    with _chat_state_lock:
        if not os.path.exists(DATA_DIR):
            os.makedirs(DATA_DIR)
        tmp = CHAT_STATE_FILE + ".tmp"
        with open(tmp, 'w', encoding='utf-8') as f:
            json.dump(state, f, ensure_ascii=False, indent=2)
        os.replace(tmp, CHAT_STATE_FILE)

# Инициализация состояния диалога из файла
_chat_state = _load_chat_state()

def get_chat_state(chat_id, key):
    """Возвращает значение состояния для чата (или None)."""
    return _chat_state.get(str(chat_id), {}).get(key)

def set_chat_state(chat_id, key, value):
    """Устанавливает значение состояния для чата и сохраняет на диск."""
    cid = str(chat_id)
    if cid not in _chat_state:
        _chat_state[cid] = {}
    if value is None:
        _chat_state[cid].pop(key, None)
    else:
        _chat_state[cid][key] = value
    _save_chat_state(_chat_state)

# ============================================================
#  GIT: АВТОКОММИТ И АВТОПУШ КОНФИГОВ
# ============================================================

import subprocess

# Полный путь к git (в PATH его может не быть)
_GIT_EXE = r"C:\Program Files\Git\bin\git.exe"

def _git(*args):
    """Выполняет git-команду и возвращает (returncode, output)."""
    try:
        result = subprocess.run(
            [_GIT_EXE] + list(args),
            capture_output=True, text=True, encoding="utf-8", errors="replace"
        )
        return result.returncode, (result.stdout + result.stderr).strip()
    except Exception as e:
        return 1, str(e)

def git_commit_and_push(files, message):
    """Добавляет файлы, коммитит и пушит. Возвращает (ok, text)."""
    code, out = _git("add", "--", *files)
    if code != 0:
        return False, f"git add: {out}"
    code, out = _git("commit", "-m", message)
    if code != 0 and "nothing to commit" not in out.lower():
        return False, f"git commit: {out}"
    code, out = _git("push", "origin", "master")
    if code != 0:
        return False, f"git push: {out}"
    return True, out

# ============================================================
#  ДОБАВЛЕНИЕ СУДОВ И КОМПАНИЙ (С АВТОПУШЕМ)
# ============================================================

def add_ship(name):
    """Добавляет судно в ships.json и пушит. Возвращает (ok, text)."""
    name = name.strip()
    if not name:
        return False, "Пустое название судна."
    ships = load_ships()
    key = name.lower()
    if key in ships:
        return False, f"Судно «{name}» уже есть в списке."
    ships[key] = name
    with open(SHIPS_FILE, 'w', encoding='utf-8') as f:
        json.dump(ships, f, ensure_ascii=False, indent=2)
    ok, text = git_commit_and_push([SHIPS_FILE], f"feat: добавить судно {name}")
    if ok:
        return True, f"✅ Судно «{name}» добавлено и запушено в репозиторий."
    return True, f"✅ Судно «{name}» добавлено в файл, но пуш не удался: {text}"

def add_company(field, value):
    """Обновляет поле компании (executor/customer/location) и пушит."""
    value = value.strip()
    if not value:
        return False, "Пустое значение."
    companies = load_companies()
    companies[field] = value
    with open(COMPANIES_FILE, 'w', encoding='utf-8') as f:
        json.dump(companies, f, ensure_ascii=False, indent=2)
    ok, text = git_commit_and_push([COMPANIES_FILE], f"feat: обновить {field} в companies.json")
    if ok:
        return True, f"✅ Поле «{field}» обновлено и запушено в репозиторий."
    return True, f"✅ Поле «{field}» обновлено в файле, но пуш не удался: {text}"

# ============================================================
#  ОПРЕДЕЛЕНИЕ ТИПА ОБОРУДОВАНИЯ (ЛОКАЛЬНОЕ)
# ============================================================

def detect_equipment_type(text):
    """Определяет тип оборудования локально (без AI)"""
    text_lower = text.lower()
    if any(word in text_lower for word in ["двигател", "дизель", "мотор", "гд"]):
        return "engine"
    elif any(word in text_lower for word in ["компрессор", "компрес"]):
        return "compressor"
    elif any(word in text_lower for word in ["насос", "помп"]):
        return "pump"
    else:
        return None

def ask_for_clarification(equipment):
    """Формирует вопрос для уточнения типа оборудования"""
    return f"🔍 Я нашёл упоминание '{equipment}'. Уточните, это:\n1️⃣ Насос\n2️⃣ Двигатель\n3️⃣ Другое оборудование\n\nПросто напишите номер или название."

# ============================================================
#  РАСШИРЕННЫЙ АНАЛИЗ ЗАПРОСОВ
# ============================================================

def detect_ship(text):
    text_lower = text.lower()
    ships = load_ships()
    for key, name in ships.items():
        if key in text_lower:
            return name
    return None

def detect_pump_type(text):
    text_lower = text.lower()
    
    piston_keywords = ["поршн", "плунж", "прямодейств", "паровой"]
    for kw in piston_keywords:
        if kw in text_lower:
            return "piston"
    
    gear_keywords = ["шестерен", "шестерн", "ротан", "rotan", "зубчат", "маслян"]
    for kw in gear_keywords:
        if kw in text_lower:
            return "gear"
    
    centrifugal_keywords = ["центробеж", "центр", "крыльчатк"]
    for kw in centrifugal_keywords:
        if kw in text_lower:
            return "centrifugal"
    
    if "насос" in text_lower:
        if any(kw in text_lower for kw in ["маслян", "ротан"]):
            return "gear"
        if any(kw in text_lower for kw in ["крыльчатк", "центр"]):
            return "centrifugal"
        if any(kw in text_lower for kw in ["поршн", "плунж"]):
            return "piston"
    
    return None

def extract_equipment(text):
    text_lower = text.lower()
    equipment_keywords = ["насос", "двигатель", "компрессор", "вентилятор", 
                         "генератор", "кран", "лебедка", "редуктор", "гидромотор",
                         "брашпиль", "котёл", "водонагреватель", "дизель", "мотор"]
    for kw in equipment_keywords:
        if kw in text_lower:
            pattern = r'(\w+\s+){0,2}' + kw + r'(\s+\w+){0,2}'
            match = re.search(pattern, text)
            if match:
                return match.group(0).strip()
    return None

def extract_clearances_from_text(text):
    text_lower = text.lower()
    clearances = []
    
    clearance_map = {
        "радиальн": "radial",
        "осев": "axial", 
        "подшипник": "bearing",
        "сальник": "seal",
        "цилиндр": "cylinder_piston",
        "канавк": "ring_groove",
        "замк": "ring_gap",
        "крейцкопф": "crosshead",
        "коренн": "main_bearing",
        "шатун": "connecting_rod",
        "грундбукс": "seal_wear"
    }
    
    patterns = [
        r'зазор\s+(\w+)\s+(\d+\.?\d*)',
        r'(\w+)\s+зазор\s+(\d+\.?\d*)',
        r'зазор\s+(\d+\.?\d*)\s+(\w+)',
        r'зазор\s+(\d+\.?\d*)',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            if isinstance(match, tuple):
                parts = list(match)
                value = None
                clearance_type = None
                for part in parts:
                    if re.match(r'^\d+\.?\d*$', part):
                        value = float(part)
                    elif part in clearance_map:
                        clearance_type = clearance_map[part]
                    elif part in ["radial", "axial", "bearing", "seal"]:
                        clearance_type = part
                if value is not None:
                    clearances.append({
                        "type": clearance_type or "unknown",
                        "value": value,
                        "raw": match
                    })
            else:
                if re.match(r'^\d+\.?\d*$', match):
                    for ct, mapped in clearance_map.items():
                        if ct in text_lower:
                            clearances.append({
                                "type": mapped,
                                "value": float(match),
                                "raw": match
                            })
                            break
    return clearances

def extract_defects(text):
    text_lower = text.lower()
    defects = []
    
    if "дефекты" in text_lower:
        defect_part = re.split(r'дефекты[:;]', text_lower, flags=re.IGNORECASE)
        if len(defect_part) > 1:
            parts = defect_part[1].strip().split(',')
            parts = [p.strip() for p in parts if p.strip()]
            for p in parts:
                if p:
                    defects.append(p)
            if defects:
                return defects
    
    defect_keywords = [
        "поврежден", "повреждена", "повреждено", "повреждены",
        "сгнил", "сгнила", "сгнило", "сгнили",
        "высох", "высохла", "высохло", "высохли",
        "треснул", "треснула", "треснуло", "треснули",
        "сломан", "сломана", "сломано", "сломаны",
        "разбит", "разбита", "разбито", "разбиты",
        "износ", "течь", "коррози", "трещин", "разруш", "выкрашиван", 
        "задир", "деформац", "ржав", "люфт", "биение", "стук", "вибрац",
        "зазор", "перегрев", "заедание", "отказ", "неисправн", "поломк",
        "изгиб", "скручиван", "ослаблен", "изношен", "выработк",
        "закоксовыван", "загрязнен", "неплотн", "подтекани", "протечк"
    ]
    
    found_defects = []
    sentences = re.split(r'[,.!?;]', text)
    for sentence in sentences:
        sentence_lower = sentence.lower().strip()
        if not sentence_lower:
            continue
        for kw in defect_keywords:
            if kw in sentence_lower:
                found_defects.append(sentence.strip())
                break
    
    if found_defects:
        return found_defects
    
    if "зазор" in text_lower:
        clearances = extract_clearances_from_text(text)
        for c in clearances:
            found_defects.append(f"зазор {c['type']}: {c['value']} мм")
        return found_defects
    
    return []

def parse_works_for_avr(text):
    works = []
    text_lower = text.lower()
    
    clean_text = re.sub(r'(авр|акт выполненных|сделай авр|создай авр|оформи авр)\s*', '', text_lower, flags=re.IGNORECASE)
    clean_text = re.sub(r'по судну\s+\w+\s*', '', clean_text)
    clean_text = re.sub(r'судно\s+\w+\s*', '', clean_text)
    clean_text = clean_text.strip()
    
    lines = re.split(r'\n|\.\s+|;\s+', clean_text)
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 5:
            continue
        
        work = {"name": "", "description": "", "quantity": "", "unit": "", "note": ""}
        
        quantity_match = re.search(r'(\d+)\s*(шт|компл|м|кг|л|шт\.|компл\.|м\.|кг\.|л\.)', line)
        if quantity_match:
            work["quantity"] = quantity_match.group(1)
            work["unit"] = quantity_match.group(2).replace('.', '')
            line = line.replace(quantity_match.group(0), '').strip()
        
        note_match = re.search(r'\([^)]+\)', line)
        if note_match:
            work["note"] = note_match.group(0).strip('()')
            line = line.replace(note_match.group(0), '').strip()
        
        if ':' in line or '—' in line or '-' in line:
            parts = re.split(r':\s*|—\s*|-\s*', line, maxsplit=1)
            if len(parts) == 2:
                work["name"] = parts[0].strip().capitalize()
                work["description"] = parts[1].strip().capitalize()
            else:
                work["description"] = line.capitalize()
        else:
            if any(word in line for word in ["замена", "ремонт", "восстановлен", "изготовлен", "монтаж", "демонтаж"]):
                work["name"] = "Ремонтные работы"
                work["description"] = line.capitalize()
            else:
                work["description"] = line.capitalize()
        
        if work["name"] or work["description"]:
            if not work["unit"]:
                work["unit"] = "компл." if not work["quantity"] else ""
            works.append(work)
    
    if not works and text.strip():
        works.append({
            "name": "Основные работы",
            "description": text.strip().capitalize(),
            "quantity": "1",
            "unit": "компл.",
            "note": ""
        })
    
    return works

def analyze_query_local(text):
    """Локальный анализ запроса (без AI)"""
    result = {
        "ship": detect_ship(text),
        "equipment": extract_equipment(text),
        "defects": extract_defects(text),
        "pump_type": detect_pump_type(text),
        "equipment_type": detect_equipment_type(text),
        "clearances": extract_clearances_from_text(text),
        "works": parse_works_for_avr(text),
        "full_text": text,
        "source": "local"
    }
    
    if result["clearances"] and not result["defects"]:
        for c in result["clearances"]:
            result["defects"].append(f"зазор {c['type']}: {c['value']} мм")
    
    if not result["equipment"] and "насос" in text.lower():
        pump_name = pump_db.get_pump_name(result["pump_type"]) if result["pump_type"] else ""
        result["equipment"] = f"насос {pump_name}".strip() if pump_name else "насос"
    
    return result

# ============================================================
#  ГЕНЕРАЦИЯ ОБЪЁМА РАБОТ
# ============================================================

def generate_work_volume(defects, full_text, pump_type=None, equipment_type=None):
    """Генерирует объём работ с использованием базы знаний или AI"""
    
    # 1. Пробуем Алису через ai_router
    if alisa_router:
        try:
            ai_result = alisa_router.generate_work_volume(defects, equipment_type, pump_type)
            if ai_result:
                print(f"✅ Сгенерировано через Алису")
                return ai_result
        except Exception as e:
            print(f"⚠️ Ошибка при вызове Алисы: {e}")
    
    # 2. Если ничего не помогло — базовый шаблон
    print("⚠️ Использую базовый шаблон")
    return generate_base_work_volume(defects)

def generate_base_work_volume(defects):
    """Базовый объём работ (запасной вариант)"""
    lines = ["1. Демонтаж узла", "2. Разборка и дефектация"]
    
    work_items = []
    for defect in defects:
        defect_lower = defect.lower()
        if "течь" in defect_lower or "уплотнен" in defect_lower:
            work_items.append("замена уплотнительных элементов")
        elif "износ" in defect_lower or "изношен" in defect_lower:
            work_items.append("восстановление или замена изношенных деталей")
        elif "трещин" in defect_lower:
            work_items.append("заварка трещин или замена детали")
        elif "коррози" in defect_lower:
            work_items.append("зачистка и восстановление коррозионных повреждений")
        elif "зазор" in defect_lower:
            work_items.append("регулировка зазоров")
        elif "протечк" in defect_lower:
            work_items.append("замена уплотнений и проверка герметичности")
    
    if work_items:
        unique_items = list(dict.fromkeys(work_items))
        lines.append("3. " + "; ".join(unique_items))
    else:
        lines.append("3. Замена/восстановление деталей")
    
    lines.append("4. Сборка с проверкой зазоров")
    lines.append("5. Монтаж")
    lines.append("6. Предъявление лицу сдающему")
    
    return "\n".join(lines)

# ============================================================
#  ПОСТРОЕНИЕ ТАБЛИЦЫ (ДЛЯ НАСОСОВ)
# ============================================================

DEFECT_MAP = {
    "крыльчатк": "2.1",
    "крылатк": "2.1",
    "колес": "2.1",
    "вал": "2.2",
    "трещин вала": "2.2",
    "шпонк": "2.3",
    "уплотнен": "3.1",
    "сальник": "3.1",
    "подшипник": "4.1",
    "крепеж": "6.3",
    "болт": "6.3",
    "гайк": "6.3",
}

DEFAULT_NO_DEFECT_TEXT = "Визуальный осмотр. Дефектов не обнаружено."
DEFAULT_NO_DEFECT_WORK = "Мыть, чистить. Годен к дальнейшей эксплуатации."

# Карта "ключевое слово в дефекте" -> "короткое действие"
ACTION_MAP = {
    "эксплуатационный износ": "Замена.",
    "износ": "Замена.",
    "коррози": "Чистка УШМ, грунтовка. Пригодна к дальнейшей эксплуатации.",
    "грязев": "Мыть, чистить. Годен к дальнейшей эксплуатации.",
    "окисление": "Мыть, чистить. Годен к дальнейшей эксплуатации.",
    "трещин": "Замена.",
    "течь": "Замена уплотнений, проверка герметичности.",
}

def build_defect_table_pump(pump_type, defects, work_volume):
    """Строит таблицу для насосов (7 колонок)"""
    rows = [
        {"num": "1.1", "part": "Корпус насоса", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "1.2", "part": "Уплотнительное кольцо", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "1.3", "part": "Всасывающий/напорный патрубок", "defect": "", "unit": "компл.", "qty": "1"},
    ]
    
    if pump_type == "centrifugal":
        rows.extend([
            {"num": "2.1", "part": "Рабочее колесо (крыльчатка)", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.2", "part": "Вал насоса", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.3", "part": "Шпонка рабочего колеса", "defect": "", "unit": "шт.", "qty": "1"},
        ])
    elif pump_type == "gear":
        rows.extend([
            {"num": "2.1", "part": "Ведущая шестерня", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.2", "part": "Ведомая шестерня", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.3", "part": "Пальцы и втулки", "defect": "", "unit": "компл.", "qty": "1"},
            {"num": "2.4", "part": "Перепускной клапан", "defect": "", "unit": "шт.", "qty": "1"},
        ])
    elif pump_type == "piston":
        rows.extend([
            {"num": "2.1", "part": "Цилиндр (зеркало)", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.2", "part": "Поршень / плунжер", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.3", "part": "Поршневые кольца", "defect": "", "unit": "компл.", "qty": "1"},
            {"num": "2.4", "part": "Шток / плунжер", "defect": "", "unit": "шт.", "qty": "1"},
            {"num": "2.5", "part": "Крейцкопф (башмаки, направляющие)", "defect": "", "unit": "компл.", "qty": "1"},
        ])
    
    rows.extend([
        {"num": "3.1", "part": "Уплотнение вала (сальник/торцевое)", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "4.1", "part": "Подшипниковый узел", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "4.2", "part": "Корпус подшипников / крышки", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "4.3", "part": "Масляная камера", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "5.1", "part": "Обмотка статора", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "5.2", "part": "Клеммная коробка и кабельный ввод", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "5.3", "part": "Муфта соединения", "defect": "", "unit": "компл.", "qty": "1"},
        {"num": "6.1", "part": "Обратный клапан", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "6.2", "part": "Задвижка на всасывании", "defect": "", "unit": "шт.", "qty": "1"},
        {"num": "6.3", "part": "Крепёж и расходные материалы", "defect": "", "unit": "компл.", "qty": "1"},
    ])
    
    for defect in defects:
        defect_lower = defect.lower()
        placed = False
        for keyword, pos in DEFECT_MAP.items():
            if keyword in defect_lower:
                for row in rows:
                    if row["num"] == pos and not row["defect"]:
                        row["defect"] = defect
                        placed = True
                        break
                if placed:
                    break
        if not placed:
            for row in rows:
                if row["num"] == "1.1" and not row["defect"]:
                    row["defect"] = defect
                    break
    
    for row in rows:
        if not row["defect"]:
            row["defect"] = DEFAULT_NO_DEFECT_TEXT
            row["work"] = DEFAULT_NO_DEFECT_WORK
        else:
            defect_lower = row["defect"].lower()
            matched = None
            for keyword, action in ACTION_MAP.items():
                if keyword in defect_lower:
                    matched = action
                    break
            row["work"] = matched or work_volume
    
    return rows

# ============================================================
#  ПОСТРОЕНИЕ ТАБЛИЦЫ (ДЛЯ ДВИГАТЕЛЕЙ)
# ============================================================

def build_defect_table_engine(defects, work_volume):
    """Строит таблицу для двигателей (6 колонок)"""
    rows = []
    
    sections = {
        "Цилиндропоршневая группа": [],
        "Головка цилиндров и газораспределение": [],
        "Системы и вспомогательное оборудование": [],
        "Сборочные и испытательные работы": []
    }
    
    for i, defect in enumerate(defects, 1):
        defect_lower = defect.lower()
        section = None
        
        if any(word in defect_lower for word in ["поршн", "цилиндр", "кольц", "втулк"]):
            section = "Цилиндропоршневая группа"
        elif any(word in defect_lower for word in ["крышк", "клапан", "форсунк", "толкател", "газораспредел"]):
            section = "Головка цилиндров и газораспределение"
        elif any(word in defect_lower for word in ["рубашк", "сальник", "турбо", "масл", "охлажд"]):
            section = "Системы и вспомогательное оборудование"
        elif any(word in defect_lower for word in ["испытани", "сборк"]):
            section = "Сборочные и испытательные работы"
        else:
            section = "Прочее"
        
        if not defect or defect == "Не указано":
            defect_text = DEFAULT_NO_DEFECT_TEXT
            work_text = DEFAULT_NO_DEFECT_WORK
        else:
            defect_lower = defect.lower()
            matched = None
            for keyword, action in ACTION_MAP.items():
                if keyword in defect_lower:
                    matched = action
                    break
            defect_text = defect
            work_text = matched or work_volume

        rows.append({
            "num": str(i),
            "defect": defect_text,
            "work": work_text,
            "unit": "компл.",
            "qty": "1",
            "section": section
        })
    
    return rows

# ============================================================
#  ФУНКЦИИ ОРМ: НАГРУЖКА РЕМОНТНОЙ ВЕДОМОСТИ
# ============================================================

def save_repair_items_to_db(ship_id, items):
    """
    Сохраняет пункты ремонтной ведомости в БД с дедупликацией.
    Возвращает: (inserted_count, skipped_count, statement_id)
    """
    session = SessionLocal()
    try:
        # Создать запись в repair_statements
        stmt = RepairStatement(ship_id=ship_id, source_excel_file_ref="uploaded")
        session.add(stmt)
        session.flush()  # Получить ID
        statement_id = stmt.id
        
        inserted = 0
        skipped = 0
        
        for item in items:
            # Проверить дубликат (по item_number + section)
            existing = session.query(StatementItem).filter(
                StatementItem.statement_id == statement_id,
                StatementItem.item_number == item.get("item_number"),
                StatementItem.section == item.get("section")
            ).first()
            
            if existing:
                skipped += 1
                continue
            
            stmt_item = StatementItem(
                statement_id=statement_id,
                item_number=item.get("item_number"),
                description=item.get("description"),
                quantity=item.get("quantity"),
                section=item.get("section"),
                status="active"
            )
            session.add(stmt_item)
            inserted += 1
        
        session.commit()
        return inserted, skipped, statement_id
    except Exception as e:
        session.rollback()
        raise e
    finally:
        session.close()


def get_user_role(telegram_id):
    """Получить роль пользователя из ORM."""
    session = SessionLocal()
    try:
        user = session.query(User).filter_by(telegram_id=telegram_id).first()
        return user.role if user else "customer"
    finally:
        session.close()


def can_upload_repair_list(telegram_id):
    """Проверить, может ли пользователь загружать ремонтную ведомость.
    Доступно всем, кроме customer.
    """
    role = get_user_role(telegram_id)
    return role != "customer"


# ============================================================
#  ФУНКЦИИ ВЕРСИОНИРОВАНИЯ ДОКУМЕНТОВ
# ============================================================

def count_drafts_for_item(item_id, category="defect_act_draft"):
    """Кол-во draft'ов для данного item_id и категории."""
    session = SessionLocal()
    try:
        count = session.query(Document).filter(
            Document.item_id == item_id,
            Document.category == category,
            Document.status == "draft"
        ).count()
        return count
    finally:
        session.close()


def get_oldest_draft(item_id, category="defect_act_draft"):
    """Получить старейший draft для item_id."""
    session = SessionLocal()
    try:
        doc = session.query(Document).filter(
            Document.item_id == item_id,
            Document.category == category,
            Document.status == "draft"
        ).order_by(Document.uploaded_at.asc()).first()
        return doc
    finally:
        session.close()


def handle_document_approve(document_id, user_id):
    """
    Утвердить документ: draft → approved.
    Возвращает: (success, message)
    """
    session = SessionLocal()
    try:
        doc = session.query(Document).filter_by(id=document_id).first()
        if not doc:
            return False, "❌ Документ не найден"
        
        if doc.status != "draft":
            return False, f"❌ Документ уже {doc.status}"
        
        doc.status = "approved"
        session.commit()
        return True, f"✅ Документ утверждён"
    except Exception as e:
        session.rollback()
        return False, f"❌ Ошибка: {str(e)}"
    finally:
        session.close()


def handle_document_archive(document_id, user_id):
    """
    Архивировать документ: approved → archived.
    Только ADMIN_IDS.
    Возвращает: (success, message)
    """
    if user_id not in ADMIN_IDS:
        return False, "🚫 Только админы могут архивировать документы"
    
    session = SessionLocal()
    try:
        doc = session.query(Document).filter_by(id=document_id).first()
        if not doc:
            return False, "❌ Документ не найден"
        
        if doc.status != "approved":
            return False, f"❌ Можно архивировать только approved документы (текущий: {doc.status})"
        
        doc.status = "archived"
        session.commit()
        return True, f"✅ Документ архивирован"
    except Exception as e:
        session.rollback()
        return False, f"❌ Ошибка: {str(e)}"
    finally:
        session.close()


def handle_document_delete(document_id, user_id):
    """
    Удалить документ.
    - draft: любой может удалить
    - approved: только ADMIN_IDS
    Возвращает: (success, message)
    """
    session = SessionLocal()
    try:
        doc = session.query(Document).filter_by(id=document_id).first()
        if not doc:
            return False, "❌ Документ не найден"
        
        if doc.status == "approved" and user_id not in ADMIN_IDS:
            return False, "🚫 Только админы могут удалять approved документы"
        
        # Удалить файл (FileStorage проверит статус)
        if doc.file_ref:
            storage.delete_file(doc.file_ref)
        
        # Удалить из БД
        session.delete(doc)
        session.commit()
        return True, f"✅ Документ удалён"
    except Exception as e:
        session.rollback()
        return False, f"❌ Ошибка: {str(e)}"
    finally:
        session.close()


# ============================================================
#  ФУНКЦИИ СОЗДАНИЯ ДОКУМЕНТОВ
# ============================================================

def create_defect_document(ship, equipment, defects, work_volume, pump_type=None, repair_type=None, purpose=None, basis=None):
    """Создаёт акт дефектации с таблицей, подходящей под тип оборудования"""
    doc = load_template("defect_act_template.docx")
    
    number = get_next_number("da")
    
    date_str = datetime.now().strftime('%d.%m.%Y')
    ship_code = ship[:3].upper() if ship else "XXX"
    
    equipment_type = detect_equipment_type(equipment or "")
    if equipment_type is None:
        equipment_type = "pump"
    
    if equipment_type == "pump":
        rows_data = build_defect_table_pump(pump_type, defects, work_volume)
        cols = 7
        headers = ['№', 'Позиция', 'Дефект / Состояние', 'Объём работ', 'Ед. изм.', 'Кол-во', 'Примечание']
        sections = {
            "1": "Корпус и проточная часть",
            "2": "Ротор / рабочая часть",
            "3": "Уплотнения вала",
            "4": "Подшипниковый узел",
            "5": "Электропривод",
            "6": "Арматура и обвязка"
        }
        get_section_key = lambda row: row["num"].split(".")[0]
        show_purpose = True
        show_basis = True
        show_conclusion = True
        show_notes = False
        notes_text = ""
    else:
        rows_data = build_defect_table_engine(defects, work_volume)
        cols = 6
        headers = ['№ п/п', 'Наименование дефекта', 'Объём работ', 'Ед. изм.', 'Кол-во', 'Примечание']
        sections = {}
        get_section_key = lambda row: row.get("section", "Прочее")
        show_purpose = False
        show_basis = False
        show_conclusion = False
        show_notes = True
        notes_text = "Все СЗЧ (поршневые кольца, поршни, втулка, комплекты для форсунок, РТИ) — поставка Заказчика, если не указано иное.\nРаботы по проточке и транспортировке деталей выполняются Подрядчиком за отдельную плату (акт дополнительных работ)."
    
    table_paragraph_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "{{table}}" in paragraph.text:
            table_paragraph_index = i
            paragraph.text = ""
            break
    
    table = doc.add_table(rows=1, cols=cols)
    table.autofit = False
    table.allow_autofit = False
    
    if cols == 7:
        widths = [Cm(1.3), Cm(3.8), Cm(5.0), Cm(5.0), Cm(2.0), Cm(1.8), Cm(3.8)]
    else:
        widths = [Cm(1.8), Cm(5.0), Cm(6.0), Cm(2.2), Cm(2.0), Cm(3.0)]
    
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    current_section = None
    for row_data in rows_data:
        section_key = get_section_key(row_data)
        if section_key != current_section:
            current_section = section_key
            row = table.add_row().cells
            for cell in row:
                cell.text = ""
            if cols == 7:
                row[0].text = sections.get(section_key, "")
            else:
                row[0].text = section_key
            for cell in row:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        row = table.add_row().cells
        if cols == 7:
            row[0].text = row_data["num"]
            row[1].text = row_data["part"]
            row[2].text = row_data.get("defect", "—")
            row[3].text = row_data.get("work", "—")
            row[4].text = row_data["unit"]
            row[5].text = row_data["qty"]
            row[6].text = "—"
        else:
            row[0].text = row_data["num"]
            row[1].text = row_data.get("defect", "—")
            row[2].text = row_data.get("work", "—")
            row[3].text = row_data.get("unit", "компл.")
            row[4].text = row_data.get("qty", "1")
            row[5].text = "—"
    
    if table_paragraph_index is not None:
        target_paragraph = doc.paragraphs[table_paragraph_index]
        tbl = table._tbl
        target_paragraph._element.addprevious(tbl)
    
    placeholders = {
        "ship_code": ship_code,
        "number": str(number).zfill(2),
        "date": date_str,
        "ship": ship or "Не указано",
        "equipment": equipment or "Не указано",
        "work_object": repair_type or "Текущий ремонт",
        "purpose": purpose or "Определение технического состояния и объема ремонта",
        "basis": basis or f"План-график ремонта на {datetime.now().year} год"
    }
    
    if show_conclusion:
        placeholders["conclusion"] = "Детали подлежат замене/восстановлению согласно указанному объёму работ."
    placeholders["special_notes"] = notes_text if show_notes else ""
    
    doc = replace_placeholders(doc, placeholders)
    
    if equipment_type != "pump":
        for paragraph in doc.paragraphs:
            if "Представитель подрядчика (Исполнитель)" in paragraph.text:
                paragraph.text = "Представитель Подрядчика:"
            if "Представитель заказчика (Судовладелец / Экипаж)" in paragraph.text:
                paragraph.text = "Представитель Заказчика:"
            if "Старший механик" in paragraph.text:
                paragraph.text = "Должность      / *[ФИО]* /"
            if "Согласовано (при необходимости)" in paragraph.text:
                paragraph.text = ""
            if "Инспектор РМРС" in paragraph.text:
                paragraph.text = ""
    
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_avr_document(ship, works, executor=None, customer=None, location=None):
    companies = load_companies()
    executor = executor or companies.get("executor")
    customer = customer or companies.get("customer")
    location = location or companies.get("location")
    doc = load_template("avr_template.docx")
    
    number = get_next_number("avr")
    
    date_str = datetime.now().strftime('%d.%m.%Y')
    ship_code = ship[:3].upper() if ship else "XXX"
    
    table_paragraph_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "{{table}}" in paragraph.text:
            table_paragraph_index = i
            paragraph.text = ""
            break
    
    table = doc.add_table(rows=1, cols=6)
    table.autofit = False
    table.allow_autofit = False
    
    widths = [Cm(1.8), Cm(5.0), Cm(6.0), Cm(2.2), Cm(2.0), Cm(3.0)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    headers = ['№ п/п', 'Наименование работ', 'Описание выполненных работ', 'Кол-во', 'Ед. изм.', 'Примечание']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if works:
        for i, work in enumerate(works, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = work.get('name', '')
            row[2].text = work.get('description', '')
            row[3].text = str(work.get('quantity', ''))
            row[4].text = work.get('unit', '')
            row[5].text = work.get('note', '')
    else:
        row = table.add_row().cells
        row[0].text = "1"
        row[1].text = "Основные работы"
        row[2].text = "Выполнены работы согласно дефектации"
        row[3].text = "1"
        row[4].text = "компл."
        row[5].text = ""
    
    if table_paragraph_index is not None:
        target_paragraph = doc.paragraphs[table_paragraph_index]
        tbl = table._tbl
        target_paragraph._element.addprevious(tbl)
    
    placeholders = {
        "ship_code": ship_code,
        "number": str(number).zfill(2),
        "date": date_str,
        "ship": ship or "Не указано",
        "executor": executor,
        "customer": customer,
        "location": location,
    }
    
    doc = replace_placeholders(doc, placeholders)
    
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ============================================================
#  КОМАНДА /START
# ============================================================

@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.reply_to(message, 
        "👋 Привет! Я — твой инженерный ассистент.\n\n"
        "📌 Что я умею:\n"
        "• Создавать Акты дефектации (скажи 'сделай акт')\n"
        "• Создавать Акты выполненных работ (скажи 'сделай АВР')\n"
        "• Проверять зазоры по ТУ (скажи 'проверь зазор')\n"
        "• Проверять параметры по ГОСТам (скажи 'проверь по ГОСТ')\n"
        "• Показывать частые дефекты (спроси 'какие дефекты')\n"
        "• Показывать чек-лист деталей (спроси 'чек-лист насоса')\n\n"
        "📌 Типы оборудования в базе:\n"
        "• Насосы: центробежные, шестерёнчатые, поршневые\n"
        "• Двигатели (MAN, Caterpillar и др.)\n\n"
        "📌 Доступные команды:\n"
        "• /gosts — список всех ГОСТов\n"
        "• /search — поиск по ГОСТам\n"
        "• /stats — статистика AI\n\n"
        "🧠 Я использую Яндекс.Алису и базу знаний для анализа запросов!\n\n"
        "📝 Примеры:\n"
        "• 'Судно Славянская, пожарный насос, повреждена крылатка. Сделай акт'\n"
        "• 'Судно Аргака, главный двигатель MAN, износ поршневых колец. Сделай акт'\n"
        "• 'проверь по ГОСТ 520-2011 диаметр=50'\n"
        "• 'проверь по ГОСТ 3325-85 зазор=0.15'"
    )

# ============================================================
#  АВТОРИЗАЦИЯ И РОЛИ
# ============================================================

@bot.message_handler(commands=['login'])
def cmd_login(message):
    """Регистрация/вход пользователя."""
    user = db.get_user(message.chat.id)
    if user and user.get("approved"):
        bot.reply_to(message, f"✅ Вы уже авторизованы как {user['name']} ({db.ROLE_LABELS.get(user['role'], user['role'])}).")
        return
    if user and not user.get("approved"):
        bot.reply_to(message, "⏳ Ваша заявка ещё на рассмотрении. Ожидайте одобрения.")
        return
    # Начинаем регистрацию: спрашиваем ФИО
    set_chat_state(message.chat.id, "reg_step", "name")
    bot.reply_to(message, "📝 Регистрация. Введите ваше ФИО:")


@bot.message_handler(commands=['approve'])
def cmd_approve(message):
    """Одобрение/отклонение заявок на регистрацию (инженер-технолог или директор)."""
    user = db.get_user(message.chat.id)
    if not db.can_approve_users(user):
        bot.reply_to(message, "🚫 У вас нет прав на одобрение пользователей.")
        return
    pending = db.get_pending_users()
    if not pending:
        bot.reply_to(message, "📭 Нет заявок на одобрение.")
        return
    lines = ["📋 Заявки на регистрацию:"]
    for p in pending:
        lines.append(f"{p['user_id']}: {p['name']} — {db.ROLE_LABELS.get(p['role_requested'], p['role_requested'])}")
    lines.append("\nОтветьте: /approve_yes <id> или /approve_no <id>")
    bot.reply_to(message, "\n".join(lines))


@bot.message_handler(commands=['approve_yes'])
def cmd_approve_yes(message):
    user = db.get_user(message.chat.id)
    if not db.can_approve_users(user):
        bot.reply_to(message, "🚫 Нет прав.")
        return
    parts = message.text.split()
    if len(parts) < 2:
        bot.reply_to(message, "Укажите id: /approve_yes <id>")
        return
    try:
        uid = int(parts[1])
    except ValueError:
        bot.reply_to(message, "Неверный id.")
        return
    pending = db.get_pending_users()
    target = next((p for p in pending if p['user_id'] == uid), None)
    if not target:
        bot.reply_to(message, "Заявка не найдена.")
        return
    db.create_user(uid, target['name'], target['role_requested'], target.get('phone'), approved=1)
    db.remove_pending_user(uid)
    db.log_action(user['user_id'], "approve_user", details=f"Одобрен пользователь {target['name']} ({uid})")
    bot.reply_to(message, f"✅ Пользователь {target['name']} одобрен.")
    try:
        bot.send_message(uid, f"✅ Ваша регистрация одобрена. Добро пожаловать, {target['name']}!")
    except Exception:
        pass


@bot.message_handler(commands=['approve_no'])
def cmd_approve_no(message):
    user = db.get_user(message.chat.id)
    if not db.can_approve_users(user):
        bot.reply_to(message, "🚫 Нет прав.")
        return
    parts = message.text.split()
    if len(parts) < 2:
        bot.reply_to(message, "Укажите id: /approve_no <id>")
        return
    try:
        uid = int(parts[1])
    except ValueError:
        bot.reply_to(message, "Неверный id.")
        return
    db.remove_pending_user(uid)
    bot.reply_to(message, f"❌ Заявка {uid} отклонена.")
    try:
        bot.send_message(uid, "❌ Ваша заявка на регистрацию отклонена.")
    except Exception:
        pass


@bot.message_handler(commands=['users'])
def cmd_users(message):
    """Список пользователей (для инженера-технолога)."""
    user = db.get_user(message.chat.id)
    if not db.is_engineer(user):
        bot.reply_to(message, "🚫 Только для инженера-технолога.")
        return
    conn = db._connect()
    cur = conn.cursor()
    cur.execute("SELECT * FROM users ORDER BY name")
    rows = cur.fetchall()
    conn.close()
    if not rows:
        bot.reply_to(message, "Пользователей пока нет.")
        return
    lines = ["👥 Пользователи:"]
    for r in rows:
        status = "✅" if r["approved"] else "⏳"
        lines.append(f"{status} {r['name']} — {db.ROLE_LABELS.get(r['role'], r['role'])}")
    bot.reply_to(message, "\n".join(lines))


# ============================================================
#  УСТАНОВКА РОЛИ (ADMIN)
# ============================================================

@bot.message_handler(commands=['set_role'])
def cmd_set_role(message):
    """Устанавливает роль пользователю: /set_role <telegram_id> <role>.
    Доступно только админам (ADMIN_IDS). Роли: technologist, user.
    """
    if message.chat.id not in ADMIN_IDS:
        bot.reply_to(message, "🚫 Команда доступна только администраторам.")
        return
    parts = message.text.split()
    if len(parts) != 3:
        bot.reply_to(message, "📝 Использование: /set_role <telegram_id> <role>\nРоли: technologist, user")
        return
    try:
        tg_id = int(parts[1])
    except ValueError:
        bot.reply_to(message, "❌ telegram_id должен быть числом.")
        return
    role = parts[2].strip().lower()
    if role not in ("technologist", "user"):
        bot.reply_to(message, "❌ Неизвестная роль. Допустимые: technologist, user")
        return
    # Сохраняем роль в таблицу users (новая схема: users.telegram_id)
    from models import SessionLocal, User
    session = SessionLocal()
    try:
        user = session.query(User).filter(User.telegram_id == tg_id).first()
        if user:
            user.role = role
        else:
            session.add(User(telegram_id=tg_id, role=role))
        session.commit()
        bot.reply_to(message, f"✅ Роль пользователя {tg_id} установлена: {role}")
    except Exception as e:
        session.rollback()
        bot.reply_to(message, f"❌ Ошибка при сохранении роли: {e}")
    finally:
        session.close()


# ============================================================
#  НОВЫЕ ОБРАБОТЧИКИ: ДОКУМЕНТООБОРОТ И НАВИГАЦИЯ
# ============================================================

if DOCUMENT_MANAGER_AVAILABLE:
    # Регистрируем обработчики загрузки ремонтной ведомости
    bot_handlers_new.register_upload_handlers(bot)
    # Регистрируем обработчики навигации
    bot_handlers_new.register_navigation_handlers(bot)


# ============================================================
#  СКАНИРОВАНИЕ ПАПКИ repair_docs
# ============================================================

@bot.message_handler(commands=['scan'])
def cmd_scan(message):
    """Сканирует папку repair_docs и обрабатывает новые файлы."""
    user = db.get_user(message.chat.id)
    if not user or not user.get("approved"):
        bot.reply_to(message, "🔒 Сначала авторизуйтесь: /login")
        return
    bot.reply_to(message, "🔍 Сканирую папку repair_docs...")
    import scanner
    messages = scanner.scan_repair_docs(uploaded_by=user["user_id"])
    for m in messages:
        bot.send_message(message.chat.id, m)
    # Уведомляем инженера/директоров о договорах на утверждение
    notify_contracts_for_approval()


@bot.message_handler(content_types=['document'])
def handle_repair_list_upload(message):
    """
    Обработчик загрузки Excel-файла с ремонтной ведомостью.
    Доступно всем, кроме customer.
    """
    if not can_upload_repair_list(message.chat.id):
        bot.reply_to(message, "🚫 У вас нет прав на загрузку ремонтной ведомости.")
        return
    
    file_info = bot.get_file(message.document.file_id)
    file_path = file_info.file_path
    downloaded_file = bot.download_file(file_path)
    
    # Сохранить временный файл
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(downloaded_file)
        tmp_path = tmp.name
    
    try:
        import scanner
        items = scanner.parse_repair_list(tmp_path)
        if not items:
            bot.reply_to(message, "⚠️ В файле не найдено пунктов ремонтной ведомости.")
            return
        
        # Получить судно из имени файла или спросить
        filename = message.document.file_name or "unknown"
        ship_name = scanner.detect_ship_from_filename(filename)
        
        if not ship_name:
            bot.reply_to(message, "❌ Не удалось определить судно из имени файла. Используйте формат: Ремведомость_<Судно>.xlsx")
            return
        
        # Найти судно в БД
        session = SessionLocal()
        try:
            ship = session.query(Ship).filter_by(name=ship_name).first()
            if not ship:
                bot.reply_to(message, f"❌ Судно '{ship_name}' не найдено в базе. Добавьте его сначала.")
                return
            
            # Сохранить пункты в БД
            inserted, skipped, stmt_id = save_repair_items_to_db(ship.id, items)
            bot.reply_to(message, 
                f"✅ Ремонтная ведомость загружена для судна '{ship_name}'\n"
                f"📝 Добавлено: {inserted} пунктов\n"
                f"⏭️ Пропущено (дубликаты): {skipped}")
        finally:
            session.close()
    except Exception as e:
        bot.reply_to(message, f"❌ Ошибка при обработке файла: {str(e)}")
    finally:
        import os
        os.unlink(tmp_path)


def notify_contracts_for_approval():
    """Уведомляет инженера-технолога и директоров о договорах, ожидающих утверждения."""
    conn = db._connect()
    cur = conn.cursor()
    cur.execute(
        "SELECT d.doc_id, s.name AS ship_name FROM documents d "
        "JOIN ships s ON s.ship_id = d.ship_id "
        "WHERE d.doc_type = ? AND d.approved = 0",
        (db.DOC_CONTRACT,),
    )
    pending = cur.fetchall()
    cur.execute(
        "SELECT user_id FROM users WHERE role IN (?, ?) AND approved = 1",
        (db.ROLE_ENGINEER, db.ROLE_DIRECTOR),
    )
    approvers = [r["user_id"] for r in cur.fetchall()]
    conn.close()
    if not pending:
        return
    for uid in approvers:
        try:
            lines = ["📄 Договоры, ожидающие утверждения:"]
            for p in pending:
                lines.append(f"• {p['ship_name']} (id={p['doc_id']})")
            lines.append("\nОтветьте: /approve_contract <id> или /reject_contract <id>")
            bot.send_message(uid, "\n".join(lines))
        except Exception:
            pass


@bot.message_handler(commands=['approve_contract'])
def cmd_approve_contract(message):
    """Утверждение договора (инженер-технолог или директор)."""
    user = db.get_user(message.chat.id)
    if not db.can_approve_users(user):
        bot.reply_to(message, "🚫 Нет прав на утверждение договоров.")
        return
    parts = message.text.split()
    if len(parts) < 2:
        bot.reply_to(message, "Укажите id: /approve_contract <id>")
        return
    try:
        doc_id = int(parts[1])
    except ValueError:
        bot.reply_to(message, "Неверный id.")
        return
    doc = db.get_document(doc_id)
    if not doc or doc["doc_type"] != db.DOC_CONTRACT:
        bot.reply_to(message, "Договор не найден.")
        return
    db.approve_document(doc_id)
    db.log_action(user["user_id"], "approve_contract", ship_id=doc["ship_id"], doc_id=doc_id)
    bot.reply_to(message, f"✅ Договор (id={doc_id}) утверждён.")


@bot.message_handler(commands=['reject_contract'])
def cmd_reject_contract(message):
    """Отклонение договора (инженер-технолог или директор)."""
    user = db.get_user(message.chat.id)
    if not db.can_approve_users(user):
        bot.reply_to(message, "🚫 Нет прав.")
        return
    parts = message.text.split()
    if len(parts) < 2:
        bot.reply_to(message, "Укажите id: /reject_contract <id>")
        return
    try:
        doc_id = int(parts[1])
    except ValueError:
        bot.reply_to(message, "Неверный id.")
        return
    doc = db.get_document(doc_id)
    if not doc:
        bot.reply_to(message, "Договор не найден.")
        return
    db.delete_document(doc_id)
    db.log_action(user["user_id"], "reject_contract", ship_id=doc["ship_id"], doc_id=doc_id)
    bot.reply_to(message, f"❌ Договор (id={doc_id}) отклонён и удалён.")


@bot.message_handler(commands=['stats'])
def show_stats(message):
    """Показывает статистику использования AI"""
    if alisa_router:
        try:
            stats = alisa_router.get_stats()
            response = "📊 **Статистика Алисы (YandexGPT):**\n\n"
            response += f"✅ Вызовов: {stats['calls']}\n"
            response += f"❌ Ошибок: {stats['errors']}\n"
            bot.reply_to(message, response, parse_mode='Markdown')
            return
        except Exception as e:
            print(f"⚠️ Ошибка при получении статистики: {e}")
    
    bot.reply_to(message, "❌ Статистика недоступна")

# ============================================================
#  КОМАНДА /GOSTS — СПИСОК ВСЕХ ГОСТОВ
# ============================================================

@bot.message_handler(commands=['gosts'])
def show_gosts(message):
    """Показывает список всех доступных ГОСТов"""
    if not gost_checker:
        bot.reply_to(message, "❌ ГОСТ чекер не загружен. Проверьте файл gost_checker.py")
        return
    
    gosts = gost_checker.get_all_gosts()
    if not gosts:
        bot.reply_to(message, "❌ База ГОСТов не загружена. Запустите merge_gosts.py")
        return
    
    response = "📁 **Доступные ГОСТы:**\n\n"
    
    # Группируем по разделам
    sections = {}
    for gost_id, data in gosts.items():
        section = data.get("section", "Общие")
        if section not in sections:
            sections[section] = []
        sections[section].append((gost_id, data.get("title", "")[:50]))
    
    for section, items in sections.items():
        response += f"**{section}** ({len(items)})\n"
        for gost_id, title in items[:5]:
            response += f"• {gost_id} — {title}...\n"
        if len(items) > 5:
            response += f"  _... и ещё {len(items)-5}_\n"
        response += "\n"
    
    response += "💡 Используйте `проверь по ГОСТ {номер} {параметр}={значение}`\n"
    response += "Пример: `проверь по ГОСТ 520-2011 диаметр=50`"
    
    bot.reply_to(message, response, parse_mode='Markdown')

# ============================================================
#  КОМАНДА /SEARCH — ПОИСК ПО ГОСТАМ
# ============================================================

@bot.message_handler(commands=['search'])
def search_gosts(message):
    """Поиск по ГОСТам"""
    if not gost_checker:
        bot.reply_to(message, "❌ ГОСТ чекер не загружен")
        return
    
    query = message.text.replace('/search', '').strip()
    if not query:
        bot.reply_to(message, "📝 Введите поисковый запрос: `/search подшипник`")
        return
    
    results = gost_checker.search(query)
    if not results:
        bot.reply_to(message, f"❌ Ничего не найдено по запросу '{query}'")
        return
    
    response = f"📋 **Результаты поиска по '{query}':**\n\n"
    for gost_id, data in list(results.items())[:10]:
        response += f"• **{gost_id}** — {data.get('title', 'Без названия')[:60]}...\n"
    
    if len(results) > 10:
        response += f"\n_... и ещё {len(results)-10} результатов_"
    
    bot.reply_to(message, response, parse_mode='Markdown')

# ============================================================
#  ОБРАБОТЧИКИ CALLBACK'OB (НАВИГАЦИЯ)
# ============================================================

@bot.callback_query_handler(func=lambda call: call.data.startswith('section_'))
def handle_section_callback(call):
    """Обработчик выбора раздела."""
    parts = call.data.split('_')
    if len(parts) < 3:
        return
    
    ship_id = int(parts[1])
    section_hash = parts[2]
    
    # Найти раздел по хешу (все разделы для судна)
    sections = navigation.get_sections_for_ship(ship_id)
    section = None
    for s in sections:
        if str(hash(s) & 0x7fffffff) == section_hash:
            section = s
            break
    
    if not section:
        bot.answer_callback_query(call.id, "❌ Раздел не найден")
        return
    
    # Показать пункты в разделе
    keyboard = navigation.build_items_keyboard(ship_id, section, page=0)
    if not keyboard:
        bot.answer_callback_query(call.id, "⚠️ В этом разделе нет пунктов")
        return
    
    text = f"📄 **Раздел:** {section}\n\nВыберите пункт:"
    bot.edit_message_text(text, call.message.chat.id, call.message.message_id, 
                         reply_markup=keyboard, parse_mode='Markdown')
    bot.answer_callback_query(call.id)


@bot.callback_query_handler(func=lambda call: call.data.startswith('item_'))
def handle_item_callback(call):
    """Обработчик выбора пункта."""
    parts = call.data.split('_')
    if len(parts) < 2:
        return
    
    item_id = int(parts[1])
    item = navigation.get_item_details(item_id)
    
    if not item:
        bot.answer_callback_query(call.id, "❌ Пункт не найден")
        return
    
    text = navigation.format_item_details(item)
    keyboard = types.InlineKeyboardMarkup()
    keyboard.add(types.InlineKeyboardButton("📄 Загрузить документ", callback_data=f"upload_doc_{item_id}"))
    keyboard.add(types.InlineKeyboardButton("⬅️ Назад", callback_data="back_to_sections"))
    
    bot.edit_message_text(text, call.message.chat.id, call.message.message_id,
                         reply_markup=keyboard, parse_mode='Markdown')
    bot.answer_callback_query(call.id)


# ============================================================
#  ГЛАВНЫЙ ОБРАБОТЧИК (ЧЕРЕЗ АЛИСУ)
# ============================================================

# История диалогов для контекста (для Алисы)
user_histories = {}

@bot.message_handler(func=lambda message: True)
def handle_message(message):
    user_text = message.text
    text_lower = user_text.lower()
    
    if user_text.startswith('/'):
        return
    
    # --- ПРОВЕРКА РОЛИ (интеграция с NLP) ---
    # Только engineer_technologist и админы могут использовать NLP-режим
    role = get_user_role(message.chat.id)
    if message.chat.id not in ADMIN_IDS and role != "engineer_technologist":
        bot.reply_to(message, "📄 Отправьте документы или используйте кнопки для навигации.")
        return
    
    # --- РЕГИСТРАЦИЯ НОВОГО ПОЛьЗОВАТЕЛЯ ---
    reg_step = get_chat_state(message.chat.id, "reg_step")
    if reg_step:
        if reg_step == "name":
            name = user_text.strip()
            if not name:
                bot.reply_to(message, "Введите ФИО:")
                return
            set_chat_state(message.chat.id, "reg_name", name)
            set_chat_state(message.chat.id, "reg_step", "role")
            bot.reply_to(
                message,
                "Выберите роль:\n"
                "1️⃣ Строитель\n"
                "2️⃣ Директор\n"
                "3️⃣ Заказчик\n\n"
                "Напишите номер или название. Если вы инженер-технолог — введите секретный код."
            )
            return
        if reg_step == "role":
            name = get_chat_state(message.chat.id, "reg_name")
            choice = user_text.strip().lower()
            # Секретный код инженера-технолога
            if ENGINEER_CODE and choice == ENGINEER_CODE.lower():
                db.create_user(message.chat.id, name, db.ROLE_ENGINEER, approved=1)
                set_chat_state(message.chat.id, "reg_step", None)
                set_chat_state(message.chat.id, "reg_name", None)
                bot.reply_to(message, f"✅ Вы зарегистрированы как инженер-технолог, {name}!")
                return
            role_map = {
                "1": db.ROLE_BUILDER, "строитель": db.ROLE_BUILDER,
                "2": db.ROLE_DIRECTOR, "директор": db.ROLE_DIRECTOR,
                "3": db.ROLE_CUSTOMER, "заказчик": db.ROLE_CUSTOMER,
            }
            role = role_map.get(choice)
            if not role:
                bot.reply_to(message, "Неверный выбор. Напишите номер (1/2/3) или название роли.")
                return
            # Инженер-технолог и директор требуют одобрения
            db.add_pending_user(message.chat.id, name, role)
            set_chat_state(message.chat.id, "reg_step", None)
            set_chat_state(message.chat.id, "reg_name", None)
            bot.reply_to(
                message,
                f"📝 Заявка на роль «{db.ROLE_LABELS.get(role, role)}» отправлена на одобрение.\n"
                "Ожидайте подтверждения."
            )
            # Уведомляем инженера-технолога и директоров
            conn = db._connect()
            cur = conn.cursor()
            cur.execute("SELECT user_id FROM users WHERE role IN (?, ?) AND approved = 1",
                        (db.ROLE_ENGINEER, db.ROLE_DIRECTOR))
            approvers = [r["user_id"] for r in cur.fetchall()]
            conn.close()
            for uid in approvers:
                try:
                    bot.send_message(uid, f"📋 Новая заявка: {name} ({db.ROLE_LABELS.get(role, role)}). /approve")
                except Exception:
                    pass
            return
    
    # --- ПРОВЕРКА АВТОРИЗАЦИИ ---
    user = db.get_user(message.chat.id)
    if not user or not user.get("approved"):
        bot.reply_to(
            message,
            "🔒 Для работы с ботом необходимо авторизоваться.\n"
            "Введите /login для регистрации или входа."
        )
        return

    # --- ПРОВЕРКА РОЛИ (НОВОЕ) ---
    # Если пользователь customer → показать меню
    if DOCUMENT_MANAGER_AVAILABLE:
        user_role = document_manager.get_user_role(message.chat.id)
        if user_role == document_manager.ROLE_CUSTOMER:
            from telebot import types
            markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
            markup.add("📋 Ремонтная ведомость")
            markup.add("📄 Документы")
            markup.add("🚢 Суда")
            bot.send_message(
                message.chat.id,
                "👋 Используйте кнопки для навигации.",
                reply_markup=markup
            )
            return

    # --- ОБРАБОТКА УТОЧНЕНИЯ ОСНОВАНИЯ АКТА ---
    pending = get_chat_state(message.chat.id, "pending_act")
    if pending:
        if text_lower.strip() in ("отмена", "отменить", "cancel", "стоп"):
            set_chat_state(message.chat.id, "pending_act", None)
            bot.reply_to(message, "Отменено. Акт не создан.")
            return
        basis = user_text.strip()
        if not basis:
            bot.reply_to(message, "Пожалуйста, укажите основание для акта.")
            return
        set_chat_state(message.chat.id, "pending_act", None)
        try:
            file_stream = create_defect_document(
                pending["ship"], pending["equipment"], pending["defects"],
                pending["work_volume"], pending["pump_type"], pending["repair_type"],
                purpose=pending["purpose"], basis=basis
            )
            bot.send_document(
                message.chat.id,
                file_stream,
                visible_file_name=f'Акт_дефектации_{pending["ship"] or "судна"}.docx'
            )
            bot.send_message(message.chat.id, "📄 Акт дефектации в Word отправлен!")
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            bot.send_message(message.chat.id, f"❌ Ошибка при создании акта:\n\n{str(e)}")
        return

    # --- ОБРАБОТКА ВЫБОРА/ДОБАВЛЕНИЯ СУДНА ---
    if get_chat_state(message.chat.id, "awaiting_ship"):
        ships = load_ships()
        ship_names = list(ships.values())
        choice = user_text.strip()
        # Выбор по номеру из списка
        if choice.isdigit():
            idx = int(choice) - 1
            if 0 <= idx < len(ship_names):
                ship = ship_names[idx]
                set_chat_state(message.chat.id, "awaiting_ship", None)
                set_chat_state(message.chat.id, "pending_act", {
                    "ship": ship,
                    "equipment": get_chat_state(message.chat.id, "draft_equipment"),
                    "defects": get_chat_state(message.chat.id, "draft_defects"),
                    "work_volume": get_chat_state(message.chat.id, "draft_work_volume"),
                    "pump_type": get_chat_state(message.chat.id, "draft_pump_type"),
                    "repair_type": get_chat_state(message.chat.id, "draft_repair_type"),
                    "purpose": "Определение технического состояния и объема ремонта",
                })
                bot.send_message(
                    message.chat.id,
                    f"✅ Судно: {ship}. Укажите основание для акта, например:\n"
                    "«План-график ремонта на 2026 год» или «Заявка капитана»"
                )
                return
            else:
                bot.reply_to(message, "❌ Неверный номер. Выберите из списка или напишите название нового судна.")
                return
        # Добавление нового судна
        if choice.lower() in ("новое", "добавить", "новое судно"):
            set_chat_state(message.chat.id, "awaiting_ship", "new")
            bot.reply_to(message, "✏️ Напишите название нового судна:")
            return
        # Если в состоянии "new" — это название нового судна
        if get_chat_state(message.chat.id, "awaiting_ship") == "new":
            ok, text = add_ship(choice)
            if ok:
                set_chat_state(message.chat.id, "awaiting_ship", None)
                set_chat_state(message.chat.id, "pending_act", {
                    "ship": choice.strip(),
                    "equipment": get_chat_state(message.chat.id, "draft_equipment"),
                    "defects": get_chat_state(message.chat.id, "draft_defects"),
                    "work_volume": get_chat_state(message.chat.id, "draft_work_volume"),
                    "pump_type": get_chat_state(message.chat.id, "draft_pump_type"),
                    "repair_type": get_chat_state(message.chat.id, "draft_repair_type"),
                    "purpose": "Определение технического состояния и объема ремонта",
                })
                bot.send_message(message.chat.id, text)
                bot.send_message(
                    message.chat.id,
                    "Укажите основание для акта, например:\n"
                    "«План-график ремонта на 2026 год» или «Заявка капитана»"
                )
            else:
                bot.reply_to(message, text)
            return
        # Прямое название судна (не номер)
        if choice:
            ok, text = add_ship(choice)
            if ok:
                set_chat_state(message.chat.id, "awaiting_ship", None)
                set_chat_state(message.chat.id, "pending_act", {
                    "ship": choice.strip(),
                    "equipment": get_chat_state(message.chat.id, "draft_equipment"),
                    "defects": get_chat_state(message.chat.id, "draft_defects"),
                    "work_volume": get_chat_state(message.chat.id, "draft_work_volume"),
                    "pump_type": get_chat_state(message.chat.id, "draft_pump_type"),
                    "repair_type": get_chat_state(message.chat.id, "draft_repair_type"),
                    "purpose": "Определение технического состояния и объема ремонта",
                })
                bot.send_message(message.chat.id, text)
                bot.send_message(
                    message.chat.id,
                    "Укажите основание для акта, например:\n"
                    "«План-график ремонта на 2026 год» или «Заявка капитана»"
                )
            else:
                bot.reply_to(message, text)
            return

    # --- ОБРАБОТКА УТОЧНЕНИЙ ---
    if get_chat_state(message.chat.id, "clarification"):
        equipment_type = text_lower
        if "1" in equipment_type or "насос" in equipment_type:
            set_chat_state(message.chat.id, "clarification", "pump")
            bot.reply_to(message, "✅ Принято: насос")
            return
        elif "2" in equipment_type or "двигател" in equipment_type:
            set_chat_state(message.chat.id, "clarification", "engine")
            bot.reply_to(message, "✅ Принято: двигатель")
            return
        else:
            set_chat_state(message.chat.id, "clarification", "other")
            bot.reply_to(message, "✅ Принято: другое оборудование")
            return
    
    # ---- 1. АКТ ДЕФЕКТАЦИИ (ЧЕРЕЗ АЛИСУ) ----
    if any(word in text_lower for word in ['сделай акт', 'акт дефектации', 'оформи акт', 'составь акт']):
        handle_act_creation(message, user_text)
        return
    
    # ---- 2. АВР ----
    if any(word in text_lower for word in ['авр', 'акт выполненных', 'сделай авр', 'оформи авр']):
        handle_avr_creation(message, user_text)
        return
    
    # ---- 3. ПРОВЕРКА ПО ГОСТАМ ----
    if any(word in text_lower for word in ['проверь по госту', 'по ГОСТ', 'по госту', 'гост']):
        gost_match = re.search(r'гост\s*([\d-]+)', user_text, re.IGNORECASE)
        if gost_match and gost_checker:
            gost_id = gost_match.group(1)
            param_match = re.search(r'(\w+)\s*[=:]\s*([\d.]+)', user_text)
            if param_match:
                param_name = param_match.group(1).strip()
                value = float(param_match.group(2))
                result = gost_checker.check_parameter(gost_id, param_name, value)
                
                response = f"📊 **Проверка по ГОСТ {gost_id}**\n\n"
                response += f"🔹 Параметр: {param_name}\n"
                response += f"🔹 Значение: {value}\n\n"
                response += f"{result.get('message', '')}"
                if result.get('action'):
                    response += f"\n\n🔧 **Рекомендация:** {result['action']}"
                
                bot.reply_to(message, response, parse_mode='Markdown')
                return
    
    # ---- 4. ВСЁ ОСТАЛЬНОЕ — ЧЕРЕЗ АЛИСУ ----
    if alisa_router:
        user_id = message.chat.id
        if user_id not in user_histories:
            user_histories[user_id] = []
        
        history = user_histories[user_id]
        
        try:
            bot.send_chat_action(message.chat.id, 'typing')
            result = alisa_router.process_query(user_text, history)
            
            history.append(f"Пользователь: {user_text}")
            history.append(f"Бот: {result.get('response', '')[:200]}")
            if len(history) > 10:
                history = history[-10:]
            user_histories[user_id] = history
            
            if result.get('status') == 'ok':
                bot.reply_to(message, result.get('response', 'Извините, не удалось получить ответ.'))
            else:
                bot.reply_to(message, "🤔 Попробую ответить без Алисы...")
                handle_local_fallback(message, user_text)
                
        except Exception as e:
            print(f"⚠️ Ошибка при вызове Алисы: {e}")
            bot.reply_to(message, "⚠️ Произошла ошибка при обращении к Алисе. Отвечаю в локальном режиме.")
            handle_local_fallback(message, user_text)
    else:
        handle_local_fallback(message, user_text)


# ============================================================
#  ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ ОБРАБОТЧИКА
# ============================================================

def handle_act_creation(message, user_text):
    """Создание Акта дефектации через Алису"""
    try:
        bot.send_message(message.chat.id, "🧠 Анализирую запрос и генерирую акт через Алису...")
        
        # ---- ГЕНЕРАЦИЯ ЧЕРЕЗ АЛИСУ ----
        if alisa_act_creator:
            try:
                act_data = alisa_act_creator.generate_act_data(user_text)
                print(f"✅ Акт сгенерирован через Алису: {act_data}")
            except Exception as e:
                print(f"⚠️ Ошибка при вызове Алисы для акта: {e}")
                act_data = None
        else:
            act_data = None
        
        # ---- ЗАПАСНОЙ ВАРИАНТ — ЛОКАЛЬНЫЙ ----
        if not act_data:
            print("⚠️ Использую локальный парсер для акта")
            analysis = analyze_query_local(user_text)
            act_data = {
                "ship": analysis.get('ship') or "Не указано",
                "equipment": analysis.get('equipment') or "Не указано",
                "repair_type": "Текущий ремонт",
                "defects": analysis.get('defects', ["Не указано"]),
                "work_volume": generate_work_volume(
                    analysis.get('defects', []), 
                    user_text, 
                    analysis.get('pump_type'), 
                    analysis.get('equipment_type')
                ),
                "conclusion": "Детали подлежат замене/восстановлению согласно указанному объёму работ."
            }
        
        ship = act_data.get('ship', "Не указано")
        equipment = act_data.get('equipment', "Не указано")
        defects = act_data.get('defects', ["Не указано"])
        work_volume = act_data.get('work_volume', generate_base_work_volume(["Не указано"]))

        if not ship or ship == "Не указано":
            # Сохраняем черновик данных акта и предлагаем выбрать/добавить судно
            set_chat_state(message.chat.id, "draft_equipment", equipment)
            set_chat_state(message.chat.id, "draft_defects", defects)
            set_chat_state(message.chat.id, "draft_work_volume", work_volume)
            set_chat_state(message.chat.id, "draft_pump_type", detect_pump_type(user_text))
            set_chat_state(message.chat.id, "draft_repair_type", act_data.get('repair_type'))
            ships = load_ships()
            ship_names = list(ships.values())
            list_text = "\n".join(f"{i+1}. {name}" for i, name in enumerate(ship_names))
            set_chat_state(message.chat.id, "awaiting_ship", "choose")
            bot.send_message(
                message.chat.id,
                "🚢 Не удалось определить судно. Выберите из списка или добавьте новое:\n\n"
                f"{list_text}\n\n"
                "Напишите номер судна, название нового судна, или «новое» для добавления."
            )
            return

        if not equipment or equipment == "Не указано":
            bot.send_message(
                message.chat.id,
                "🚫 Не удалось определить оборудование. Укажите тип и модель явно."
            )
            return

        # Определяем тип оборудования для выбора таблицы
        equipment_type = detect_equipment_type(equipment or "")
        if equipment_type is None:
            equipment_type = "pump"
        
        # Определяем тип насоса
        pump_type = detect_pump_type(user_text)
        
        # Сохраняем данные акта и уточняем основание
        repair_type = act_data.get('repair_type')
        set_chat_state(message.chat.id, "pending_act", {
            "ship": ship,
            "equipment": equipment,
            "defects": defects,
            "work_volume": work_volume,
            "pump_type": pump_type,
            "repair_type": repair_type,
            "purpose": "Определение технического состояния и объема ремонта",
        })
        bot.send_message(
            message.chat.id,
            "📋 Данные акта определены. Укажите основание для акта, например:\n"
            "«План-график ремонта на 2026 год» или «Заявка капитана»"
        )
        
    except Exception as e:
        import traceback
        error_text = f"❌ Ошибка при создании акта:\n\n{str(e)}"
        bot.send_message(message.chat.id, error_text)
        print(traceback.format_exc())

def handle_avr_creation(message, user_text):
    """Создание Акта выполненных работ"""
    ship = detect_ship(user_text)
    works = parse_works_for_avr(user_text)
    
    if not works:
        bot.reply_to(message,
            "🤔 Для создания АВР опишите выполненные работы:\n"
            "Пример: 'АВР: Кабель-трасса: замена уголков 44 шт, болтов 194 шт.'"
        )
        return
    
    file_stream = create_avr_document(ship, works)
    bot.send_document(
        message.chat.id,
        file_stream,
        visible_file_name=f'АВР_{ship or "судна"}.docx'
    )
    bot.send_message(message.chat.id, "📄 Акт выполненных работ отправлен!")

def handle_local_fallback(message, user_text):
    """Локальный режим работы (без Алисы)"""
    text_lower = user_text.lower()
    
    # Проверка зазоров
    if any(word in text_lower for word in ['проверь зазор', 'проверка зазора', 'какой зазор', 'норма зазора']):
        clearances = extract_clearances_from_text(user_text)
        if clearances:
            responses = []
            for c in clearances:
                if c['type'] != 'unknown':
                    pump_type = detect_pump_type(user_text)
                    if not pump_type:
                        if "шестерен" in text_lower or "ротан" in text_lower:
                            pump_type = "gear"
                        elif "поршн" in text_lower or "плунж" in text_lower or "паровой" in text_lower:
                            pump_type = "piston"
                        else:
                            pump_type = "centrifugal"
                    
                    result = pump_db.check_clearance(pump_type, c['type'], c['value'])
                    responses.append(f"🔹 {c['type']}: {c['value']} мм -> {result['message']}")
                    
                    # Проверка по ГОСТам
                    if gost_checker:
                        if c['type'] == 'bearing':
                            gost_result = gost_checker.check_parameter("3325-85", "clearance", c['value'])
                            if gost_result.get('status') != 'error':
                                responses.append(f"   📌 ГОСТ 3325-85: {gost_result.get('message', '')}")
                        elif c['type'] == 'axial' or c['type'] == 'radial':
                            gost_result = gost_checker.check_parameter("24643-81", "runout", c['value'])
                            if gost_result.get('status') != 'error':
                                responses.append(f"   📌 ГОСТ 24643-81: {gost_result.get('message', '')}")
            
            if responses:
                response = "📊 **Результаты проверки зазоров:**\n\n" + "\n".join(responses)
                bot.reply_to(message, response, parse_mode='Markdown')
                return
    
    # Чек-лист
    if any(word in text_lower for word in ['чек-лист', 'перечень деталей', 'какие детали']):
        pump_type = detect_pump_type(user_text)
        if pump_type:
            items = pump_db.get_checklist(pump_type)
            pump_name = pump_db.get_pump_name(pump_type)
            response = f"📋 **Чек-лист для {pump_name} насоса:**\n\n"
            for i, item in enumerate(items, 1):
                name = item.get("name") if isinstance(item, dict) else item
                response += f"{i}. {name}\n"
            bot.reply_to(message, response, parse_mode='Markdown')
        else:
            bot.reply_to(message, "📌 Уточните тип насоса: центробежный, шестерёнчатый или поршневой")
        return
    
    # Дефекты
    if any(word in text_lower for word in ['какие дефекты', 'частые дефекты', 'список дефектов', 'дефекты бывают']):
        if any(word in text_lower for word in ["двигател", "дизель", "мотор"]):
            defects = pump_db.get_common_defects("engine")
            response = f"📋 **Частые дефекты двигателей:**\n\n"
            for i, defect in enumerate(defects, 1):
                response += f"{i}. {defect}\n"
            bot.reply_to(message, response, parse_mode='Markdown')
            return
        
        pump_type = detect_pump_type(user_text)
        if pump_type:
            pump_name = pump_db.get_pump_name(pump_type)
            defects = pump_db.get_common_defects(pump_type)
            response = f"📋 **Частые дефекты {pump_name} насоса:**\n\n"
            for i, defect in enumerate(defects, 1):
                method = pump_db.get_repair_method(pump_type, defect)
                method_text = f" -> {method}" if method else ""
                response += f"{i}. {defect}{method_text}\n"
            bot.reply_to(message, response, parse_mode='Markdown')
            return
        else:
            bot.reply_to(message, "📌 Уточните тип оборудования: насос или двигатель")
            return
    
    # Нормативы
    if any(word in text_lower for word in ['норматив', 'норма', 'ту', 'техническ']):
        response = "📐 **Нормативы зазоров по ТУ**\n\n"
        for pump_type in pump_db.get_pump_types():
            pump_name = pump_db.get_pump_name(pump_type)
            response += f"**{pump_name.capitalize()} насос:**\n"
            clearances = pump_db.data.get(pump_type, {}).get("clearances", {})
            for ct, data in clearances.items():
                min_val = data.get("min", 0)
                max_val = data.get("max", 0)
                unit = data.get("unit", "мм")
                response += f"  • {ct}: {min_val}-{max_val} {unit}\n"
            response += "\n"
        bot.reply_to(message, response, parse_mode='Markdown')
        return
    
    # Если ничего не подошло
    bot.reply_to(message,
        "🤔 Я не совсем понял запрос.\n\n"
        "Что нужно?\n"
        "📄 Акт дефектации — 'сделай акт'\n"
        "📋 АВР — 'сделай АВР'\n"
        "🔧 Проверить зазор — 'проверь зазор'\n"
        "📋 Дефекты — 'какие дефекты у поршневого насоса'\n"
        "📐 Нормативы — 'нормативы зазоров'\n"
        "📋 Чек-лист — 'чек-лист центробежного насоса'\n"
        "📁 Проверка по ГОСТам — 'проверь по ГОСТ 520-2011 диаметр=50'\n"
        "📋 Список ГОСТов — '/gosts'\n"
        "🔎 Поиск по ГОСТам — '/search подшипник'"
    )

# ============================================================
#  ЗАПУСК С ПОВТОРНЫМИ ПОПЫТКАМИ
# ============================================================

def start_scan_timer():
    """Запускает периодическое сканирование папки repair_docs (раз в 12 часов)."""
    def _run():
        while True:
            time.sleep(12 * 60 * 60)  # 12 часов
            try:
                import scanner
                messages = scanner.scan_repair_docs()
                for m in messages:
                    print(f"[SCAN] {m}")
                notify_contracts_for_approval()
            except Exception as e:
                print(f"[SCAN] Ошибка: {e}")
    t = threading.Thread(target=_run, daemon=True)
    t.start()


def start_bot_with_retry():
    """Запуск бота с повторными попытками подключения"""
    max_retries = 5
    retry_delay = 10
    
    for attempt in range(max_retries):
        try:
            print(f"🔄 Попытка подключения {attempt + 1}/{max_retries}...")
            bot.infinity_polling(timeout=60, long_polling_timeout=30)
            break
        except Exception as e:
            print(f"⚠️ Ошибка: {e}")
            if attempt < max_retries - 1:
                print(f"⏳ Повтор через {retry_delay} секунд...")
                time.sleep(retry_delay)
                retry_delay += 5
            else:
                print("❌ Не удалось подключиться после всех попыток")
                raise

if __name__ == '__main__':
    print("🤖 Бот-ассистент запущен!")
    print("📌 Типы оборудования в базе: насосы (центробежные, шестерёнчатые, поршневые), двигатели")
    print("📌 Доступные функции: ДА, АВР, проверка зазоров, дефекты, нормативы, чек-лист")
    
    if alisa_router:
        print("🧠 Алиса (YandexGPT) активна — все запросы проходят через неё!")
    else:
        print("⚠️ Алиса НЕ загружена — работаю в локальном режиме")
    
    if alisa_act_creator:
        print("📄 Создатель актов через Алису загружен!")
    else:
        print("⚠️ Создатель актов через Алису НЕ загружен")
    
    # Статистика по ГОСТам
    if gost_checker:
        gosts = gost_checker.get_all_gosts()
        print(f"📚 Загружено ГОСТов: {len(gosts)}")
        if gosts:
            sections = {}
            for gost_id, data in gosts.items():
                section = data.get("section", "Общие")
                sections[section] = sections.get(section, 0) + 1
            print("📋 Разделы ГОСТов:")
            for section, count in sections.items():
                print(f"   • {section}: {count}")
    else:
        print("⚠️ ГОСТ чекер не загружен")
    
    # Инициализация ORM и синхронизация судов
    init_models()
    ships_data = load_ships()
    if ships_data:
        sync_ships_from_json(ships_data)
    
    # Регистрация команд управления документами
    document_commands.register_document_commands(
        bot, ADMIN_IDS, handle_document_approve, handle_document_archive, handle_document_delete
    )
    
    # Периодическое сканирование папки repair_docs
    start_scan_timer()

    # Запуск с повторными попытками
    start_bot_with_retry()